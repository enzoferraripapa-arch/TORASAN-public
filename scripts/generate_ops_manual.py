"""
TORASAN 操作手順書 Word 文書生成スクリプト

カラーパレット: Terracotta / Sage / Sand / Charcoal / Mustard
出力: docs/manuals/TORASAN_操作手順書.docx
"""

import os
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── docx_utils から共通ユーティリティをインポート ──
sys.path.insert(0, os.path.dirname(__file__))
from docx_utils import (
    set_cell_shading,
    add_procedure_box,
    add_command_block,
    add_note,
    setup_document,
)
from spec_constants import (
    SKILL_COUNT_UNIVERSAL,
    SKILL_COUNT_DOMAIN,
    SKILL_COUNT_TOTAL,
    KNOWLEDGE_COUNT_UNIVERSAL,
    KNOWLEDGE_COUNT_DOMAIN,
    KNOWLEDGE_COUNT_TOTAL,
    PHASE_COUNT,
    OPS_INITIAL_DATE,
    FRAMEWORK_VERSION,
    PROCESS_VERSION,
)

# ============================================================
# カラーパレット（Architecture 仕様書と統一）
# ============================================================
TERRACOTTA = "E07A5F"
SAGE = "81B29A"
SAND = "F4F1DE"
CHARCOAL = "3D405B"
MUSTARD = "F2CC8F"

FONT_JP = "游ゴシック"
FONT_EN = "Segoe UI"
FONT_MONO = "Consolas"


# ============================================================
# ヘルパー
# ============================================================
def _rgb(hex_str):
    """'E07A5F' → RGBColor"""
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def add_heading_styled(doc, text, level=1):
    """カラー付き見出しを追加"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _rgb(CHARCOAL)
        run.font.name = FONT_JP
        rPr = run._element.get_or_add_rPr()
        ea = rPr.makeelement(qn("w:rFonts"), {})
        ea.set(qn("w:eastAsia"), FONT_JP)
        rPr.append(ea)
    return h


def add_body(doc, text):
    """本文段落を追加"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = FONT_JP
        run.font.size = Pt(10.5)
    return p


def add_colored_table(doc, headers, rows, header_color=CHARCOAL,
                      alt_color="F0F4F8", col_widths=None):
    """ヘッダー色をカスタマイズ可能なテーブル"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = FONT_JP
        set_cell_shading(cell, header_color)

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = FONT_JP
            if r % 2 == 1:
                set_cell_shading(cell, alt_color)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_section_divider(doc):
    """セクション間の薄い区切り線"""
    p = doc.add_paragraph()
    run = p.add_run("_" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = _rgb(SAGE)


# ============================================================
# セクション生成関数
# ============================================================

def section_cover(doc, version_label):
    """表紙（spec doc と同じ直接構築方式）"""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TORASAN")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = _rgb(CHARCOAL)
    run.font.name = FONT_JP

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("機能安全開発プロセスフレームワーク")
    run.font.size = Pt(20)
    run.font.color.rgb = _rgb(CHARCOAL)
    run.font.name = FONT_JP

    doc.add_paragraph()

    doc_title = doc.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_title.add_run("操作手順書")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = FONT_JP

    for _ in range(4):
        doc.add_paragraph()

    # 文書情報テーブル
    info_data = [
        ("文書番号", "TORASAN-OPS-001"),
        ("版数", version_label),
        ("初版作成日", OPS_INITIAL_DATE),
        ("最終更新日", date.today().strftime("%Y-%m-%d")),
        ("フレームワーク版数", FRAMEWORK_VERSION),
        ("プロセス版数", PROCESS_VERSION),
        ("対象読者", "フレームワーク利用者・管理者"),
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = "Table Grid"
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        set_cell_shading(info_table.rows[i].cells[0], "E8EDF2")
        for cell in info_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = FONT_JP

    doc.add_page_break()


def section_revision_history(doc, revision_history):
    """改版履歴"""
    add_heading_styled(doc, "改版履歴", level=1)
    rev_rows = revision_history if revision_history else [
        ["v1.0", OPS_INITIAL_DATE, "初版作成", "TORA"],
    ]
    add_colored_table(doc,
        ["版数", "日付", "変更内容", "承認"],
        rev_rows,
        header_color=CHARCOAL,
        col_widths=[2, 3, 8, 3],
    )
    doc.add_page_break()


def section_toc(doc):
    """目次"""
    add_heading_styled(doc, "目次", level=1)
    p = doc.add_paragraph("（Word の「参考資料」→「目次の更新」で自動生成してください）")
    p.italic = True
    doc.add_page_break()


def section_01_introduction(doc):
    """1. はじめに"""
    add_heading_styled(doc, "1. はじめに", level=1)

    add_heading_styled(doc, "1.1 本書の目的", level=2)
    add_body(doc,
        "本操作手順書は、TORASAN フレームワークを使用して機能安全開発プロジェクトを運用するための"
        "具体的な手順を記述する。セッション管理からフェーズ実行、品質監査まで、"
        "日常的な操作フローを段階的に説明する。"
    )

    add_heading_styled(doc, "1.2 前提条件", level=2)
    add_body(doc, "本書の手順を実行するためには、以下の環境が必要である:")
    for p_text in [
        "Claude Code（最新版）がインストール済みであること",
        "TORASAN リポジトリがローカルにクローン済みであること",
        "Git がインストール・設定済みであること",
        "install.sh による汎用スキル配布が実行済みであること",
    ]:
        doc.add_paragraph(p_text, style="List Bullet")

    add_heading_styled(doc, "1.3 コマンド記法", level=2)
    add_body(doc, "本書で使用する記法:")
    add_colored_table(doc,
        ["記法", "意味"],
        [
            ["/command", "Claude Code 内で実行するスキルコマンド"],
            ["$ command", "ターミナル（bash）で実行するシェルコマンド"],
            ["{placeholder}", "ユーザーが適切な値に置き換える部分"],
            ["[optional]", "省略可能なパラメータ"],
        ],
        header_color=CHARCOAL,
        col_widths=[4, 12],
    )

    add_heading_styled(doc, "1.4 自然言語コマンド", level=2)
    add_body(doc,
        "CLAUDE.md のコマンドマッピングにより、自然言語での指示も受け付ける。"
    )
    add_colored_table(doc,
        ["日本語での指示", "実行されるスキル"],
        [
            ["「おつです」", "/session end"],
            ["「スキル一覧見せて」", "Web GUI (localhost:3000/skills) に誘導"],
            ["「スキル改善して」", "/skill-evolve cycle"],
            ["「最新化して」", "/claude-master scan"],
            ["「マニュアル更新して」", "python scripts/generate_manuals.py"],
            ["「プロジェクト一覧」", "/repo-manage list"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    doc.add_page_break()


def section_02_setup(doc):
    """2. 初期セットアップ"""
    add_heading_styled(doc, "2. 初期セットアップ", level=1)

    add_heading_styled(doc, "2.1 TORASAN リポジトリのクローン", level=2)
    add_procedure_box(doc, "リポジトリ取得", [
        "任意のディレクトリに移動する",
        "TORASAN リポジトリをクローンする",
    ])
    add_command_block(doc, "$ git clone {TORASAN_REPO_URL}")
    add_command_block(doc, "$ cd TORASAN")

    add_section_divider(doc)

    add_heading_styled(doc, "2.2 汎用スキルの配布", level=2)
    add_body(doc,
        f"install.sh を実行して、汎用スキル（{SKILL_COUNT_UNIVERSAL}本）と"
        f"グローバルナレッジ（{KNOWLEDGE_COUNT_UNIVERSAL}本）を ~/.claude/ に配布する。"
    )
    add_procedure_box(doc, "スキル配布", [
        "TORASAN リポジトリのルートに移動する",
        "install.sh を実行する（--dry-run で事前確認可能）",
        "配布結果を確認する",
    ])
    add_command_block(doc, "$ ./install.sh --dry-run    # 事前確認（変更なし）")
    add_command_block(doc, "$ ./install.sh              # 実行")

    add_body(doc, "配布されるもの:")
    add_colored_table(doc,
        ["種別", "配布先", "数量"],
        [
            ["汎用スキル", "~/.claude/skills/{name}/", f"{SKILL_COUNT_UNIVERSAL}本"],
            ["グローバルナレッジ", "~/.claude/knowledge/", f"{KNOWLEDGE_COUNT_UNIVERSAL}本"],
            ["マニフェスト", "~/.claude/.shared-skills-manifest.json", "1ファイル"],
        ],
        header_color=SAGE,
        col_widths=[4, 7, 3],
    )

    add_note(doc, "install.sh は何度実行しても安全です（冪等性あり）。スキル更新時にも同じコマンドで配布できます。")

    add_section_divider(doc)

    add_heading_styled(doc, "2.3 開発環境の検証", level=2)
    add_body(doc, "Claude Code 内で環境検証スキルを実行し、必要なツールが揃っているか確認する。")
    add_command_block(doc, "/env-check")
    add_body(doc, "検証されるツール:")
    for t in [
        "GCC（C コンパイラ）",
        "cppcheck（静的解析）",
        "clang-tidy（LLVM リンター）",
        "Git（バージョン管理）",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_page_break()


def section_03_session(doc):
    """3. セッション管理"""
    add_heading_styled(doc, "3. セッション管理", level=1)

    add_heading_styled(doc, "3.1 セッション開始", level=2)
    add_body(doc,
        "セッション開始時に /session start を実行する。フレームワークが自動的に "
        "コンテキストを検出し、前回の引継ぎ情報を表示する。"
    )
    add_procedure_box(doc, "セッション開始手順", [
        "Claude Code を起動する",
        "/session start を実行する",
        "表示される情報を確認し、作業を開始する",
    ])
    add_command_block(doc, "/session start")

    add_body(doc, "表示される情報:")
    add_colored_table(doc,
        ["情報", "内容", "条件"],
        [
            ["ブランチ情報", "現在のブランチ名と master との差分", "常に表示"],
            ["前回の作業", "前回の到達点と次回推奨アクション", "session_state.md 存在時"],
            ["巡回提案", "技術巡回が 7 日以上未実施の場合", "7 日超過時のみ"],
            ["未解決改善項目", "前回反省会の「未」項目", "存在時のみ"],
        ],
        header_color=SAGE,
        col_widths=[4, 7, 5],
    )

    add_note(doc, "前回と同じ作業を継続する場合は /session start --continue で簡略表示になります。")

    add_section_divider(doc)

    add_heading_styled(doc, "3.2 セッション終了", level=2)
    add_body(doc,
        "セッション終了時に /session end（または「おつです」）を実行する。"
        "反省会・引継ぎ・バックアップ・master マージが順番に実行される。"
    )
    add_procedure_box(doc, "セッション終了手順", [
        "/session end を実行する（「おつです」でも可）",
        "反省会の質問に回答する",
        "スキルフィードバックを記録する（任意）",
        "引継ぎ情報が自動保存される",
        "master マージの確認に応答する",
        "完了報告を確認する",
    ])
    add_command_block(doc, "/session end")
    add_body(doc, "または:")
    add_command_block(doc, "おつです")

    add_body(doc, "終了時に実行される処理:")
    add_colored_table(doc,
        ["Step", "処理", "出力"],
        [
            ["Step 0", "コンテキスト使用率チェック", "60% 超過時に /compact 推奨"],
            ["Step 1", "規格整合性チェック", "ASIL 誤用検出・修正"],
            ["Step 1.5", "スキルフィードバック確認", "未記録スキルの記録提案"],
            ["Step 2", "反省会（レトロスペクティブ）", "議事録保存"],
            ["Step 3", "セッション引継ぎ", "session_state.md 更新"],
            ["Step 3.5", "ワークツリー整理確認", "不要 WT の削除提案"],
            ["Step 4", "バックアップ + master マージ", "タグ作成 + マージ実行"],
        ],
        header_color=CHARCOAL,
        col_widths=[2, 5, 9],
    )

    add_section_divider(doc)

    add_heading_styled(doc, "3.3 セッション状態確認", level=2)
    add_body(doc, "現在のセッション状態を非破壊的に確認する。")
    add_command_block(doc, "/session status")

    doc.add_page_break()


def section_04_new_project(doc):
    """4. 新規プロジェクト作成"""
    add_heading_styled(doc, "4. 新規プロジェクト作成", level=1)

    add_heading_styled(doc, "4.1 作成フロー概要", level=2)
    add_body(doc,
        "TORASAN リポ内で /session start を実行し、「新規プロジェクトを始める」を選択する。"
        "対話形式でプロジェクト構成を決定し、初期化まで自動実行される。"
    )
    add_procedure_box(doc, "新規プロジェクト作成", [
        "TORASAN リポで Claude Code を起動する",
        "/session start を実行する",
        "「新規プロジェクトを始める」を選択する",
        "プロジェクト名を入力する",
        "カテゴリを選択する",
        "概要を入力する",
        "（機能安全カテゴリの場合）規格選択が実行される",
        "プロジェクトが初期化される",
    ])

    add_section_divider(doc)

    add_heading_styled(doc, "4.2 カテゴリ別テンプレート", level=2)
    add_body(doc, "全カテゴリ共通で以下のディレクトリが生成される:")
    for d in [
        ".claude/skills/, .claude/knowledge/ — スキル・ナレッジ格納",
        "docs/ — 成果物・文書格納",
        "project.json — プロジェクト SoT",
        "CLAUDE.md — プロジェクト固有指示",
    ]:
        doc.add_paragraph(d, style="List Bullet")

    add_body(doc, "上記に加え、カテゴリ固有のディレクトリとフェーズ構成が適用される:")
    add_colored_table(doc,
        ["カテゴリ", "フェーズ構成", "規格", "固有ディレクトリ"],
        [
            ["組み込み（機能安全）", f"PH-01〜PH-{PHASE_COUNT}（全 {PHASE_COUNT} フェーズ）", "ISO 26262", "src/, process_records/"],
            ["デスクトップアプリ", "要件→設計→実装→テスト→リリース", "なし", "src/, tests/"],
            ["モバイルアプリ", "要件→UI設計→実装→テスト→デプロイ", "なし", "lib/, test/, assets/"],
            ["Web 開発", "要件→設計→実装→テスト→デプロイ", "なし", "src/, public/, api/"],
            ["特許・知財", "調査→出願書類→図面→レビュー", "なし", "patents/, figures/"],
            ["研究・調査", "課題設定→調査→分析→報告", "なし", "research/, data/"],
        ],
        header_color=CHARCOAL,
        col_widths=[4, 5, 3, 4],
    )

    add_section_divider(doc)

    add_heading_styled(doc, "4.3 作成後の運用開始", level=2)
    add_procedure_box(doc, "新規 PJ での開発開始", [
        "新しいプロジェクトフォルダで Claude Code セッションを開く",
        "/session start を実行する",
        "（機能安全PJの場合）/execute-phase PH-01 からフェーズ実行を開始する",
    ])
    add_note(doc, "新規 PJ 作成は TORASAN リポで行い、実際の開発は各 PJ リポで行います。")

    doc.add_page_break()


def section_05_phase_execution(doc):
    """5. フェーズ実行（機能安全プロジェクト）"""
    add_heading_styled(doc, "5. フェーズ実行（機能安全プロジェクト）", level=1)

    add_heading_styled(doc, "5.1 V-model フェーズ一覧", level=2)
    add_body(doc,
        f"TORASAN は ISO 26262 に準拠した {PHASE_COUNT} フェーズの V-model 開発プロセスを採用する。"
        "各フェーズには入力ゲート・品質ゲートが設定されており、前提条件を満たさないと次に進めない。"
    )
    add_colored_table(doc,
        ["フェーズ", "名称", "V-model 位置", "主要成果物"],
        [
            ["PH-01", "安全計画", "左翼", "安全計画書"],
            ["PH-02", "アイテム定義", "左翼", "アイテム定義書"],
            ["PH-03", "HARA（ハザード分析・リスク評価）", "左翼", "HARA シート、安全目標"],
            ["PH-04", "機能安全コンセプト", "左翼", "FSC、FSR"],
            ["PH-05", "技術安全コンセプト", "左翼", "TSC、TSR"],
            ["PH-06", "システム設計", "左翼", "ブロック図、IF 定義"],
            ["PH-07", "HW 設計", "左翼", "回路図、部品表"],
            ["PH-08", "SW 安全要求", "左翼底", "SRS（SR/DR）"],
            ["PH-09", "SW アーキテクチャ", "左翼底", "SW 構造図"],
            ["PH-10", "SW ユニット設計", "底部", "詳細設計書"],
            ["PH-11", "SW テスト仕様", "右翼底", "テスト仕様書"],
            ["PH-12", "HW 統合テスト", "右翼", "HW テスト結果"],
            ["PH-13", "システムテスト", "右翼", "システムテスト結果"],
            ["PH-14", "安全検証", "右翼", "安全検証レポート"],
            ["PH-15", "安全アセスメント", "右翼頂", "アセスメントレポート"],
        ],
        header_color=CHARCOAL,
        col_widths=[2, 4.5, 2.5, 7],
    )

    add_section_divider(doc)

    add_heading_styled(doc, "5.2 フェーズ実行の基本手順", level=2)
    add_procedure_box(doc, "フェーズ実行", [
        "現在のフェーズ状態を確認する（localhost:3000 ダッシュボード）",
        "フェーズ実行コマンドを発行する",
        "事前チェック（前フェーズ完了、入力成果物確認）の結果を確認する",
        "フェーズ手順に従って作業を実施する",
        "成果物を作成・レビューする",
        "ゲート検証が実行される",
        "全ゲート PASS で次フェーズへ移行可能になる",
    ])
    add_command_block(doc, "/execute-phase PH-{XX}")
    add_body(doc, "例: PH-05（技術安全コンセプト）を実行する場合:")
    add_command_block(doc, "/execute-phase PH-05")

    add_section_divider(doc)

    add_heading_styled(doc, "5.3 フェーズゲート検証", level=2)
    add_body(doc,
        "各フェーズの完了時にゲート検証が自動実行される。"
        "全ゲート項目が PASS になるまで次フェーズには移行できない。"
    )
    add_colored_table(doc,
        ["ゲート種別", "検証内容", "判定基準"],
        [
            ["入力ゲート", "前フェーズ成果物の存在確認", "全入力が揃っていること"],
            ["品質ゲート", "成果物の品質基準適合", "必須セクション・整合性"],
            ["トレーサビリティゲート", "要件リンクの完全性", "孤立要件なし"],
            ["レビューゲート", "レビュー実施記録", "レビュー完了"],
        ],
        header_color=TERRACOTTA,
        col_widths=[4, 6, 6],
    )

    doc.add_page_break()


def section_06_universal_skills(doc):
    """6. 汎用スキルリファレンス"""
    add_heading_styled(doc, "6. 汎用スキルリファレンス", level=1)
    add_body(doc,
        f"全プロジェクト共通で使用できる汎用スキル（{SKILL_COUNT_UNIVERSAL}本）の詳細。"
        "install.sh で ~/.claude/skills/ に配布される。"
    )

    # -- session --
    add_heading_styled(doc, "6.1 session — セッション管理", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["start", "セッション開始（自動コンテキスト検出）"],
            ["start --continue", "セッション再開（簡略表示）"],
            ["end", "セッション終了（反省会→引継ぎ→バックアップ→マージ）"],
            ["status", "セッション状態確認（非破壊）"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    # -- execute-phase --
    add_heading_styled(doc, "6.2 execute-phase — V-model フェーズ実行", level=2)
    add_body(doc, f"V-model の各フェーズ（PH-01〜PH-{PHASE_COUNT}）を実行する。")
    add_command_block(doc, "/execute-phase PH-{XX}")

    # -- dashboard --
    add_heading_styled(doc, "6.3 dashboard — ダッシュボード", level=2)
    add_body(doc,
        "プロジェクトの進捗・成果物・SPICE レベルを表示する。"
        "Web GUI (localhost:3000) に委譲される。"
    )

    # -- health-check --
    add_heading_styled(doc, "6.4 health-check — ヘルスチェック", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["quick", "主要指標のみ簡易チェック"],
            ["full", "全項目（フェーズバランス・トレーサビリティ・SPICE 成熟度）を横断分析"],
            ["action", "問題箇所に対するアクション提案"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    # -- validate --
    add_heading_styled(doc, "6.5 validate — 統合検証", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["numbers", "数値の整合性検証"],
            ["types", "型ルール検証（T-01〜T-10）"],
            ["gates", "ゲート条件検証"],
            ["all", "全項目一括検証"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    # -- trace --
    add_heading_styled(doc, "6.6 trace — トレーサビリティ", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["check", "トレーサビリティリンクの整合性チェック"],
            ["matrix", "トレーサビリティマトリクス生成"],
            ["report", "トレーサビリティレポート出力"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )
    add_body(doc, "チェーン構造:")
    add_command_block(doc, "SG → FSR → TSR → SR → TC（前方トレース）")
    add_command_block(doc, "TC → SR → TSR → FSR → SG（後方トレース）")

    # -- commit-change --
    add_heading_styled(doc, "6.7 commit-change — コミット", level=2)
    add_body(doc, "Git コミットと project.json の changeLog 記録を同時に行う。")
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["feat", "新機能コミット"],
            ["fix", "バグ修正コミット"],
            ["docs", "ドキュメント変更コミット"],
            ["refactor", "リファクタリング（機能変更なし）"],
            ["chore", "ビルド・設定変更"],
            ["test", "テスト追加・修正"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    # -- backup --
    add_heading_styled(doc, "6.8 backup — バックアップ", level=2)
    add_body(doc, "Git タグベースのスナップショットを作成する。")
    add_command_block(doc, "/backup {label}")

    # -- skill-manage --
    add_heading_styled(doc, "6.9 skill-manage — スキル管理", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["list", "全スキル（汎用/ドメイン）の一覧表示"],
            ["audit", "スキル品質基準への適合チェック"],
            ["coverage", "プロセスに対するスキルカバレッジ分析"],
            ["scaffold {name}", "スキルテンプレートの生成"],
        ],
        header_color=MUSTARD,
        col_widths=[5, 11],
    )

    # -- skill-evolve --
    add_heading_styled(doc, "6.10 skill-evolve — スキル改善エンジン", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["feedback", "スキル実行結果のフィードバック記録（1〜5 点 + コメント）"],
            ["analyze", "ギャップ分析（改善点の特定）"],
            ["improve {name}", "スキル定義に改善を適用"],
            ["maturity", "成熟度レベル（L1〜L5）の現状確認"],
            ["cycle", "feedback → analyze → improve の一括実行"],
            ["tech-refresh", "技術リフレッシュ（依存技術の更新確認）"],
        ],
        header_color=MUSTARD,
        col_widths=[5, 11],
    )

    # -- claude-master --
    add_heading_styled(doc, "6.11 claude-master — 技術巡回", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["scan", "公式ドキュメント巡回（差分検出）"],
            ["diff", "前回巡回からの差分表示"],
            ["apply", "検出された更新をスキルに反映"],
            ["domain", "ドメイン固有の更新チェック"],
            ["report", "巡回レポート出力"],
            ["schedule", "巡回スケジュール表示"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    # -- repo-manage --
    add_heading_styled(doc, "6.12 repo-manage — リポジトリ管理", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["list", "登録済みプロジェクト一覧を表示"],
            ["register {path}", "プロジェクトをレジストリに追加"],
            ["sync {name|all}", "ドメインスキルを配布"],
            ["move {name} {new-path}", "PJ 移動 + メモリ移行"],
            ["remove {name}", "レジストリから登録解除"],
            ["status", "詳細同期状態を分析"],
            ["discover", "FS 上の未登録 PJ を検出"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    # -- 残りの汎用スキル --
    add_heading_styled(doc, "6.13 その他の汎用スキル", level=2)
    add_colored_table(doc,
        ["スキル", "説明", "主なサブコマンド"],
        [
            ["env-check", "開発環境検証", "tools / path / config / fix / all"],
            ["platform-info", "OS・WSL・アーキテクチャ情報の検出", "\u2014"],
            ["assess-spice", "Automotive SPICE 能力レベル評価", "{proc} / all"],
            ["config-audit", "構成管理監査（SUP.8 対応）", "\u2014"],
            ["update-record", "SPICE プロセス記録の更新", "\u2014"],
            ["manage-tbd", "TBD 項目の管理", "list / resolve / add"],
            ["problem-resolve", "問題解決管理", "list / add / analyze / resolve / report"],
            ["worktree-cleanup", "ワークツリー・ブランチの整理", "scan / clean / branches"],
        ],
        header_color=TERRACOTTA,
        col_widths=[3.5, 5, 7.5],
    )

    doc.add_page_break()


def section_07_domain_safety(doc):
    """7. ドメインスキル — 安全規格・要件系"""
    add_heading_styled(doc, "7. ドメインスキル — 安全規格・要件系", level=1)
    add_body(doc,
        f"ドメインスキル（{SKILL_COUNT_DOMAIN}本）は /repo-manage sync で各 PJ に配布される。"
        "このセクションでは安全規格・要件系スキル（9本）を記載する。"
    )

    add_heading_styled(doc, "7.1 select-standard — 規格選定", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["derive", "製品情報から適用規格を導出"],
            ["check", "現在の規格設定を確認"],
            ["matrix", "規格対応マトリクス表示"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "7.2 switch-standard — 規格切替", level=2)
    add_body(doc, "プロジェクトの適用規格を切り替える。")
    add_command_block(doc, "/switch-standard iso26262    # ISO 26262 モードに切替")
    add_command_block(doc, "/switch-standard iec60730    # IEC 60730 モードに切替")
    add_command_block(doc, "/switch-standard iec61508    # IEC 61508 モードに切替")

    add_heading_styled(doc, "7.3 fmea — FMEA（故障モード影響分析）", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["show", "FMEA テーブル表示"],
            ["add", "故障モードの追加"],
            ["dc", "診断カバレッジ（DC）算出"],
            ["map", "故障モード → 安全機構のマッピング"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )
    add_body(doc, "RPN パラメータ:")
    add_colored_table(doc,
        ["パラメータ", "範囲", "説明"],
        [
            ["重大度（S）", "1〜10", "故障の影響の重大さ"],
            ["発生度（O）", "1〜10", "故障の発生頻度"],
            ["検出度（D）", "1〜10", "故障の検出難易度"],
            ["RPN", "S x O x D", "リスク優先度数（高いほど対策優先）"],
        ],
        header_color=TERRACOTTA,
        col_widths=[4, 2.5, 9.5],
    )

    add_heading_styled(doc, "7.4 safety-concept — 安全コンセプト", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["fsc", "機能安全コンセプト（FSC）生成"],
            ["tsc", "技術安全コンセプト（TSC）生成"],
            ["mechanism", "安全機構の設計"],
            ["redundancy", "冗長性分析"],
            ["ftti", "FTTI（フォルトトレラント時間間隔）計算"],
            ["all", "全項目一括実行"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "7.5 safety-diag — 自己診断コード生成", level=2)
    add_body(doc, "IEC 60730 Annex H に準拠する自己診断コードを生成する。")
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["all", "全診断テスト生成"],
            ["ram", "RAM テスト"],
            ["rom", "ROM テスト（CRC）"],
            ["cpu", "CPU レジスタテスト"],
            ["clock", "クロック監視テスト"],
            ["voltage", "電圧監視テスト"],
            ["wdt", "ウォッチドッグタイマテスト"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "7.6 safety-verify — 安全検証", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["verify", "安全要求の検証実行"],
            ["case", "Safety Case 構築（GSN）"],
            ["report", "検証レポート出力"],
            ["gap", "ギャップ分析"],
            ["all", "全項目一括実行"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "7.7 srs-generate — SW 安全要求仕様書", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["derive", "TSR から SW 安全要求（SR/DR）を導出"],
            ["diagnostic", "診断要求の生成"],
            ["document", "仕様書ドキュメント出力"],
            ["review", "レビューチェックリスト生成"],
            ["all", "全項目一括実行"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "7.8 system-design — システム設計", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["block", "ブロック図生成"],
            ["interface", "インターフェース定義"],
            ["allocation", "機能割当て"],
            ["asil", "ASIL 分解"],
            ["review", "設計レビュー"],
            ["all", "全項目一括実行"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "7.9 hw-review — HW 設計レビュー", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["circuit", "回路設計レビュー"],
            ["component", "部品選定レビュー"],
            ["emc", "EMC 設計レビュー"],
            ["safety", "安全設計レビュー"],
            ["integration", "統合レビュー"],
            ["all", "全項目一括実行"],
        ],
        header_color=TERRACOTTA,
        col_widths=[5, 11],
    )

    doc.add_page_break()


def section_08_domain_design(doc):
    """8. ドメインスキル — 設計・実装系"""
    add_heading_styled(doc, "8. ドメインスキル — 設計・実装系", level=1)

    add_heading_styled(doc, "8.1 sw-design — SW 設計", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["architecture", "層構造設計（APP/BSW/HAL）"],
            ["state-machine", "状態遷移設計"],
            ["module", "モジュール分割設計"],
            ["pattern", "設計パターン適用"],
            ["detailed", "詳細設計"],
            ["review", "設計レビュー"],
            ["all", "全項目一括実行"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "8.2 driver-gen — ドライバ生成", level=2)
    add_body(doc, "MISRA C:2012 準拠の HAL/BSW ドライバコードを生成する。")
    add_body(doc, "対応ペリフェラル: wdt, adc, timer, uart, spi, dem")
    add_command_block(doc, "/driver-gen {peripheral}    # 個別ドライバ生成")
    add_command_block(doc, "/driver-gen all             # 全ドライバ一括生成")

    add_heading_styled(doc, "8.3 mcu-config — MCU 設定コード生成", level=2)
    add_body(doc, "MCU（RL78/G14 等）のペリフェラル初期化コードを生成する。")
    add_body(doc, "対応ペリフェラル: adc, wdt, timer, gpio, uart, clock, all")
    add_command_block(doc, "/mcu-config {peripheral}")

    add_heading_styled(doc, "8.4 motor-control — モータ制御", level=2)
    add_body(doc, "BLDC モータ制御コードの生成・分析・テストを行う。")
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["generate", "制御コード生成"],
            ["analyze", "制御ロジック分析"],
            ["test", "テストコード生成"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "8.5 memory-map — メモリマップ", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["show", "メモリマップ表示"],
            ["analyze", "リソース使用量分析"],
            ["update", "メモリマップ更新"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "8.6 static-analysis — 静的解析", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["all", "全ツール一括実行"],
            ["cppcheck", "cppcheck 実行"],
            ["clang-tidy", "clang-tidy 実行"],
            ["flawfinder", "flawfinder 実行"],
            ["misra", "MISRA C:2012 チェック"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "8.7 ingest — 既存資産取り込み", level=2)
    add_body(doc, "既存のソースコード・文書をフレームワークに取り込む。")

    add_heading_styled(doc, "8.8 generate-docs — 成果物ドキュメント生成", level=2)
    add_body(doc, "フェーズ成果物のドキュメントを生成する。")
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["PH-{XX}", "指定フェーズの成果物を生成"],
            ["trace-matrix", "トレーサビリティマトリクスを生成"],
            ["all", "全成果物一括生成"],
        ],
        header_color=MUSTARD,
        col_widths=[5, 11],
    )

    doc.add_page_break()


def section_09_domain_test(doc):
    """9. ドメインスキル — テスト・検証系"""
    add_heading_styled(doc, "9. ドメインスキル — テスト・検証系", level=1)

    add_heading_styled(doc, "9.1 test-design — テスト設計", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["strategy", "テスト戦略策定"],
            ["cases", "テストケース生成"],
            ["boundary", "境界値テスト設計"],
            ["equivalence", "同値分割テスト設計"],
            ["negative", "異常系テスト設計"],
            ["all", "全項目一括実行"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "9.2 test-coverage — テストカバレッジ", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["requirements", "要件カバレッジ分析"],
            ["code", "コードカバレッジ分析（MC/DC、ステートメント、ブランチ）"],
            ["gap", "カバレッジギャップ分析"],
            ["all", "全項目一括実行"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "9.3 systest-design — システムテスト設計", level=2)
    add_colored_table(doc,
        ["サブコマンド", "説明"],
        [
            ["spec", "テスト仕様書生成"],
            ["scenario", "テストシナリオ設計"],
            ["acceptance", "受入条件設計"],
            ["all", "全項目一括実行"],
        ],
        header_color=TERRACOTTA,
        col_widths=[5, 11],
    )

    doc.add_page_break()


def section_10_gui(doc):
    """10. GUI（Web アプリ）"""
    add_heading_styled(doc, "10. GUI（Web アプリ）", level=1)
    add_body(doc,
        "TORASAN は React + Fastify ベースの管理 GUI（localhost:3000）を提供する。"
        "読み取り専用の情報表示は Web アプリに委譲される。"
    )

    add_heading_styled(doc, "10.1 GUI の起動", level=2)
    add_command_block(doc, "$ cd app && npm run dev")

    add_heading_styled(doc, "10.2 ページ一覧", level=2)
    add_colored_table(doc,
        ["ページ", "URL", "内容"],
        [
            ["ダッシュボード", "/", "プロジェクト全体の進捗"],
            ["スキル一覧", "/skills", "全スキルの一覧・状態"],
            ["変更履歴", "/changelog", "project.json の変更ログ"],
            ["セッション", "/sessions", "セッション履歴"],
            ["トラッカー", "/tracker", "TBD・Issues・Ideas の追跡"],
            ["環境", "/environment", "開発環境情報"],
        ],
        header_color=CHARCOAL,
        col_widths=[3.5, 3.5, 9],
    )

    doc.add_page_break()


def section_11_skill_distribution(doc):
    """11. スキル配布・同期"""
    add_heading_styled(doc, "11. スキル配布・同期", level=1)

    add_heading_styled(doc, "11.1 汎用スキルの更新・再配布", level=2)
    add_body(doc,
        "TORASAN でスキルを修正・改善した後、install.sh を再実行して全 PJ に反映する。"
    )
    add_procedure_box(doc, "汎用スキル更新", [
        "TORASAN リポでスキルを修正する",
        "変更をコミットする",
        "install.sh を実行して配布する",
    ])
    add_command_block(doc, "$ ./install.sh")

    add_section_divider(doc)

    add_heading_styled(doc, "11.2 ドメインスキルの同期", level=2)
    add_body(doc, "登録済みプロジェクトにドメインスキルを同期する。")
    add_procedure_box(doc, "ドメインスキル同期", [
        "TORASAN リポで Claude Code を起動する",
        "/repo-manage sync を実行する",
        "同期結果を確認する",
    ])
    add_command_block(doc, "/repo-manage sync all      # 全ドメインスキルを全登録PJに同期")
    add_command_block(doc, "/repo-manage sync {name}   # 特定スキルのみ同期")

    doc.add_page_break()


def section_12_quality(doc):
    """12. 品質管理・監査"""
    add_heading_styled(doc, "12. 品質管理・監査", level=1)

    add_heading_styled(doc, "12.1 ヘルスチェック", level=2)
    add_body(doc, "プロジェクト全体の健全性を評価する。")
    add_command_block(doc, "/health-check full")
    add_body(doc, "評価項目:")
    for item in [
        "V モデルバランス（左翼 vs 右翼）",
        "トレーサビリティチェーン完全性",
        "ドキュメントカバレッジ",
        "SPICE プロセス成熟度分布",
        "ギャップ・ブロッカー特定",
        "リスクアセスメント",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_section_divider(doc)

    add_heading_styled(doc, "12.2 SPICE アセスメント", level=2)
    add_command_block(doc, "/assess-spice all          # 全 16 プロセスを評価")
    add_command_block(doc, "/assess-spice SYS.1        # 個別プロセスを評価")
    add_body(doc, "評価出力:")
    add_colored_table(doc,
        ["項目", "内容"],
        [
            ["PA 1.1 評価", "ベースプラクティス実施率 + エビデンス"],
            ["PA 2.1 評価", "パフォーマンスマネジメント GP 充足度"],
            ["PA 2.2 評価", "ワークプロダクトマネジメント GP 充足度"],
            ["能力レベル判定", "Level 0〜2 の判定 + 根拠"],
            ["改善提案", "次レベル到達のための具体的アクション"],
        ],
        header_color=SAGE,
        col_widths=[4, 12],
    )

    add_section_divider(doc)

    add_heading_styled(doc, "12.3 構成管理監査", level=2)
    add_command_block(doc, "/config-audit")
    add_body(doc,
        "Automotive SPICE SUP.8 に基づく構成管理監査を実施する。"
        "ベースライン整合性、バージョン管理状態を検証する。"
    )

    add_section_divider(doc)

    add_heading_styled(doc, "12.4 バリデーション（型ルール検証）", level=2)
    add_command_block(doc, "/validate all")
    add_body(doc, "成果物間の整合性チェック（型ルール T-01〜T-10）を実行する。")
    add_colored_table(doc,
        ["ルール", "検証内容"],
        [
            ["T-01〜T-03", "型整合性（数値型、列挙型、構造体フィールド）"],
            ["T-04〜T-06", "クロスリファレンス整合性（関数名、定数値）"],
            ["T-07〜T-08", "プロセス整合性（PROCESS.md との一致）"],
            ["T-09〜T-10", "トレーサビリティ整合性（リンク切れ検出）"],
        ],
        header_color=TERRACOTTA,
        col_widths=[3, 13],
    )

    doc.add_page_break()


def section_13_problem_management(doc):
    """13. 問題管理・変更管理"""
    add_heading_styled(doc, "13. 問題管理・変更管理", level=1)

    add_heading_styled(doc, "13.1 問題報告・解決", level=2)
    add_command_block(doc, "/problem-resolve")
    add_body(doc, "SUP.9 に基づく問題管理。報告→分析→解決→クローズのサイクルで管理する。")

    add_heading_styled(doc, "13.2 変更記録", level=2)
    add_body(doc,
        "project.json の changeLog に変更履歴が記録される。"
        "変更があった場合は既存フォーマットに従って記録する。"
    )

    add_heading_styled(doc, "13.3 プロセス記録更新", level=2)
    add_command_block(doc, "/update-record")
    add_body(doc,
        "SPICE プロセス記録（process_records/ 配下）の BP 実施状況とワークプロダクトを更新する。"
    )

    add_heading_styled(doc, "13.4 TBD 管理", level=2)
    add_command_block(doc, "/manage-tbd")
    add_colored_table(doc,
        ["操作", "説明"],
        [
            ["list", "未解決 TBD 項目の一覧と依存フェーズ"],
            ["add", "新規 TBD 項目の登録"],
            ["resolve", "TBD 項目の解決記録"],
        ],
        header_color=CHARCOAL,
        col_widths=[4, 12],
    )

    doc.add_page_break()


def section_14_backup(doc):
    """14. バックアップ・復元"""
    add_heading_styled(doc, "14. バックアップ・復元", level=1)

    add_heading_styled(doc, "14.1 バックアップ作成", level=2)
    add_body(doc, "Git タグベースのスナップショットを作成する。")
    add_command_block(doc, "/backup {label}")
    add_body(doc, "例:")
    add_command_block(doc, "/backup before-ph05")
    add_body(doc, "結果: git tag v{N}-before-ph05 が作成される。")

    add_heading_styled(doc, "14.2 復元", level=2)
    add_body(doc, "タグを指定して状態を復元する。")
    add_command_block(doc, "$ git checkout v{N}-{label}")

    add_heading_styled(doc, "14.3 master マージ", level=2)
    add_body(doc, "/session end 時に自動実行される。手動で実行する場合:")
    add_command_block(doc, "$ git checkout master")
    add_command_block(doc, '$ git merge {branch} --no-ff -m "merge: {branch} セッション成果マージ"')

    doc.add_page_break()


def section_15_worktree(doc):
    """15. ワークツリー管理"""
    add_heading_styled(doc, "15. ワークツリー管理", level=1)

    add_heading_styled(doc, "15.1 ワークツリーの仕組み", level=2)
    add_body(doc,
        "TORASAN は Git ワークツリーを使用して並列セッション実行を実現する。"
        "各ワークツリーは .claude/worktrees/{session-name}/ に作成され、"
        "独立したブランチで作業できる。"
    )

    add_heading_styled(doc, "15.2 ワークツリー整理", level=2)
    add_command_block(doc, "/worktree-cleanup")
    add_body(doc,
        "不要なワークツリー（マージ済み、到達不能）を検出し、安全に削除する。"
        "/session end の Step 3.5 でも自動的に提案される。"
    )

    doc.add_page_break()


def section_16_manual_management(doc):
    """16. マニュアル管理"""
    add_heading_styled(doc, "16. マニュアル管理", level=1)

    add_heading_styled(doc, "16.1 マニュアル再生成", level=2)
    add_body(doc,
        "フレームワークの構成（スキル追加・プロセス変更等）に変更があった場合、"
        "Word 文書を再生成して最新状態に同期する。"
    )
    add_procedure_box(doc, "マニュアル再生成手順", [
        "コンテンツ生成スクリプトを編集する",
        "統合スクリプトを実行する",
        "差分レポートを確認する",
        "変更をコミットする",
    ])

    add_section_divider(doc)

    add_heading_styled(doc, "16.2 コマンドとオプション", level=2)
    add_colored_table(doc,
        ["コマンド", "説明"],
        [
            ["python scripts/generate_manuals.py", "全文書を再生成（バージョン据え置き）"],
            ['python scripts/generate_manuals.py -m "変更内容"', "再生成 + マイナーバージョンアップ + 改版履歴記録"],
            ['python scripts/generate_manuals.py -m "変更内容" --bump major', "メジャーバージョンアップ（v1.x → v2.0）"],
            ["python scripts/generate_manuals.py --spec", "仕様書のみ再生成"],
            ["python scripts/generate_manuals.py --ops", "操作手順書のみ再生成"],
            ["python scripts/generate_manuals.py --diff-only", "差分レポートのみ（再生成なし）"],
        ],
        header_color=CHARCOAL,
        col_widths=[9, 7],
    )

    add_section_divider(doc)

    add_heading_styled(doc, "16.3 バージョン管理ルール", level=2)
    add_colored_table(doc,
        ["パターン", "動作", "用途"],
        [
            ["-m なし", "再生成のみ（バージョン据え置き）", "体裁修正・リビルド"],
            ['-m "内容"', "マイナーバージョンアップ（v1.1 → v1.2）", "通常の内容変更"],
            ["--bump major", "メジャーバージョンアップ（v1.x → v2.0）", "大規模構成変更"],
        ],
        header_color=SAGE,
        col_widths=[4, 6, 6],
    )
    add_body(doc, "改版履歴・バージョンは .manuals_manifest.json で一元管理される。")

    add_section_divider(doc)

    add_heading_styled(doc, "16.4 差分レポート", level=2)
    add_body(doc,
        "再生成時に差分が検出されると、docs/manuals/diff_reports/ に "
        "unified diff 形式のレポートが保存される。"
    )
    add_body(doc, "レポートの読み方:")
    for item in [
        "+ で始まる行: 追加された内容",
        "- で始まる行: 削除された内容",
        "@@ で始まる行: 変更セクションの位置情報",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_section_divider(doc)

    add_heading_styled(doc, "16.5 運用ルール", level=2)
    add_colored_table(doc,
        ["ルール", "説明"],
        [
            ["内容変更時は必ず再生成", "docx を直接編集しない。スクリプトを修正して再生成する"],
            ["差分レポートを Git 管理", "diff_reports/ もコミットして変更履歴を残す"],
            ["マニフェストをコミット", ".manuals_manifest.json もバージョン管理する"],
            ["セッション終了時に確認", "構成変更があった場合は /session end 前に再生成する"],
        ],
        header_color=MUSTARD,
        col_widths=[5, 11],
    )
    add_note(doc, "直接 Word ファイルを編集すると、次回の再生成で上書きされます。必ずスクリプト経由で更新してください。")

    doc.add_page_break()


def section_17_troubleshooting(doc):
    """17. トラブルシューティング"""
    add_heading_styled(doc, "17. トラブルシューティング", level=1)

    add_colored_table(doc,
        ["症状", "原因", "対処法"],
        [
            [
                "/session start でプロジェクトが見つからない",
                "project.json が存在しない or パスが異なる",
                "カレントディレクトリを確認。project.json の存在を確認する",
            ],
            [
                "スキルが認識されない",
                "install.sh 未実行 or スキルファイル破損",
                "./install.sh を再実行する",
            ],
            [
                "フェーズゲートが FAIL になる",
                "前フェーズ未完了 or 成果物不足",
                "ダッシュボードで進捗確認、不足成果物を作成する",
            ],
            [
                "master マージで衝突",
                "並列作業による変更競合",
                "Git の衝突解決手順に従って手動解決する",
            ],
            [
                "技術巡回で差分が大量に出る",
                "長期間巡回していなかった",
                "差分を 1 つずつレビューし、必要なものだけ適用する",
            ],
            [
                "env-check でツールが見つからない",
                "PATH 未設定 or ツール未インストール",
                "必要なツールをインストールし PATH を通す",
            ],
            [
                "ドメインスキル同期が失敗する",
                "レジストリのパスが無効",
                "/repo-manage status で確認、必要ならパスを更新する",
            ],
            [
                "GUI (localhost:3000) が起動しない",
                "ポート競合 or npm 未インストール",
                "ポート使用状況を確認、cd app && npm install を実行",
            ],
        ],
        header_color=TERRACOTTA,
        col_widths=[5, 4.5, 6.5],
    )

    doc.add_page_break()


def section_18_periodic_tasks(doc):
    """18. 定期運用タスク"""
    add_heading_styled(doc, "18. 定期運用タスク", level=1)

    add_colored_table(doc,
        ["頻度", "タスク", "コマンド"],
        [
            ["毎セッション開始時", "セッション開始", "/session start"],
            ["毎セッション終了時", "セッション終了", "/session end"],
            ["7 日超巡回なし時", "技術巡回", "/claude-master scan"],
            ["スキル実行後", "フィードバック記録", "/skill-evolve feedback"],
            ["5 セッション毎", "スキル監査", "/skill-manage audit"],
            ['構成変更時', "マニュアル更新", 'python scripts/generate_manuals.py -m "内容"'],
        ],
        header_color=CHARCOAL,
        col_widths=[4, 4, 8],
    )

    doc.add_page_break()


def section_appendix_a_quickref(doc):
    """付録A: コマンドクイックリファレンス"""
    add_heading_styled(doc, "付録A: コマンドクイックリファレンス", level=1)

    add_heading_styled(doc, "セッション管理", level=2)
    add_colored_table(doc,
        ["コマンド", "説明"],
        [
            ["/session start", "セッション開始（自動コンテキスト検出）"],
            ["/session start --continue", "セッション再開（簡略表示）"],
            ["/session end", "セッション終了（反省会→引継ぎ→バックアップ→マージ）"],
            ["/session status", "セッション状態確認（非破壊）"],
            ["おつです", "/session end と同義"],
        ],
        header_color=CHARCOAL,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "開発プロセス", level=2)
    add_colored_table(doc,
        ["コマンド", "説明"],
        [
            ["/execute-phase PH-{XX}", "指定フェーズを実行"],
            ["/validate all", "成果物整合性検証"],
            ["/trace check", "トレーサビリティ検証"],
            ["/manage-tbd list", "TBD 項目一覧"],
        ],
        header_color=SAGE,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "品質・安全", level=2)
    add_colored_table(doc,
        ["コマンド", "説明"],
        [
            ["/assess-spice {proc|all}", "SPICE 能力レベル評価"],
            ["/health-check full", "プロジェクト健全性評価"],
            ["/config-audit", "構成管理監査（SUP.8）"],
            ["/static-analysis {target}", "静的解析実行"],
            ["/fmea show", "FMEA テーブル表示"],
            ["/safety-diag all", "診断カバレッジ設計"],
            ["/safety-concept all", "安全コンセプト設計"],
            ["/safety-verify all", "安全ケース検証（GSN）"],
        ],
        header_color=TERRACOTTA,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "スキル・フレームワーク管理", level=2)
    add_colored_table(doc,
        ["コマンド", "説明"],
        [
            ["/skill-manage list", "スキル一覧表示"],
            ["/skill-manage audit", "スキル品質監査"],
            ["/skill-evolve cycle", "スキル改善 PDCA サイクル"],
            ["/repo-manage list", "登録 PJ 一覧"],
            ["/repo-manage sync all", "ドメインスキル同期"],
            ["/claude-master scan", "技術巡回実行"],
            ["/env-check", "開発環境検証"],
        ],
        header_color=MUSTARD,
        col_widths=[5, 11],
    )

    add_heading_styled(doc, "マニュアル生成", level=2)
    add_colored_table(doc,
        ["コマンド", "説明"],
        [
            ["python scripts/generate_manuals.py", "全文書再生成"],
            ['python scripts/generate_manuals.py -m "内容"', "バージョンアップ + 再生成"],
            ["python scripts/generate_manuals.py --diff-only", "差分確認のみ"],
        ],
        header_color=CHARCOAL,
        col_widths=[8, 8],
    )

    doc.add_page_break()


def section_appendix_b_directory(doc):
    """付録B: ディレクトリ構成"""
    add_heading_styled(doc, "付録B: ディレクトリ構成", level=1)

    add_body(doc, "TORASAN リポジトリの主要ディレクトリ構成:")
    add_colored_table(doc,
        ["パス", "役割"],
        [
            [".claude/skills/", f"全スキル（汎用 {SKILL_COUNT_UNIVERSAL} + ドメイン {SKILL_COUNT_DOMAIN} = {SKILL_COUNT_TOTAL}本）"],
            [".claude/knowledge/", f"ナレッジ（汎用 {KNOWLEDGE_COUNT_UNIVERSAL} + ドメイン {KNOWLEDGE_COUNT_DOMAIN} = {KNOWLEDGE_COUNT_TOTAL}本）"],
            ["app/", "管理 GUI（React + Fastify, localhost:3000）"],
            ["docs/", "成果物ドキュメント"],
            ["docs/manuals/", "生成マニュアル（DOCX）"],
            ["docs/diagrams/", "アーキテクチャ図"],
            ["projects/", "個別PJスナップショット"],
            ["scripts/", "ユーティリティスクリプト"],
            ["install.sh", "汎用スキル配布スクリプト"],
            ["CLAUDE.md", "プロジェクト設定"],
            ["PROCESS.md", "ISO 26262 + SPICE プロセス定義"],
        ],
        header_color=CHARCOAL,
        col_widths=[4, 12],
    )

    add_body(doc,
        f"フレームワーク版数: {FRAMEWORK_VERSION} / "
        f"プロセス版数: {PROCESS_VERSION} / "
        f"スキル総数: {SKILL_COUNT_TOTAL} / "
        f"ナレッジ総数: {KNOWLEDGE_COUNT_TOTAL}"
    )


# ============================================================
# メイン
# ============================================================
def create_ops_manual(output_path, version_label='v1.0', revision_history=None):
    doc = Document()
    setup_document(doc)

    # 表紙
    section_cover(doc, version_label)

    # 改版履歴・目次
    section_revision_history(doc, revision_history)
    section_toc(doc)

    # 本文セクション
    section_01_introduction(doc)
    section_02_setup(doc)
    section_03_session(doc)
    section_04_new_project(doc)
    section_05_phase_execution(doc)
    section_06_universal_skills(doc)
    section_07_domain_safety(doc)
    section_08_domain_design(doc)
    section_09_domain_test(doc)
    section_10_gui(doc)
    section_11_skill_distribution(doc)
    section_12_quality(doc)
    section_13_problem_management(doc)
    section_14_backup(doc)
    section_15_worktree(doc)
    section_16_manual_management(doc)
    section_17_troubleshooting(doc)
    section_18_periodic_tasks(doc)
    section_appendix_a_quickref(doc)
    section_appendix_b_directory(doc)

    # 保存
    doc.save(output_path)
    print(f"操作手順書を保存しました: {output_path}")


if __name__ == "__main__":
    from pathlib import Path as _Path
    repo_root = _Path(__file__).resolve().parent.parent
    output = repo_root / "docs" / "manuals" / "TORASAN_操作手順書.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    create_ops_manual(str(output))
