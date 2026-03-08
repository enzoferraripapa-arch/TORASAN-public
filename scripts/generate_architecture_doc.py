"""
TORASAN アーキテクチャ仕様書 Word 文書生成スクリプト

カラーパレット: Terracotta / Sage / Sand / Charcoal / Mustard
出力: docs/TORASAN_Architecture.docx
"""

import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Excalidraw PNG ダイアグラムパス ──
DIAGRAMS_PNG_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)),
                                 "docs", "diagrams", "png")

# ── docx_utils から共通ユーティリティをインポート ──
sys.path.insert(0, os.path.dirname(__file__))
from docx_utils import (
    set_cell_shading,
    add_styled_table,
    add_procedure_box,
    add_command_block,
    add_note,
    setup_document,
    add_cover_page,
)
from spec_constants import (
    SKILL_COUNT_UNIVERSAL,
    SKILL_COUNT_DOMAIN,
    SKILL_COUNT_TOTAL,
    KNOWLEDGE_COUNT_UNIVERSAL,
    KNOWLEDGE_COUNT_DOMAIN,
    KNOWLEDGE_COUNT_TOTAL,
    PHASE_COUNT,
    FRAMEWORK_VERSION,
)

# ============================================================
# カラーパレット（PPTX と統一）
# ============================================================
TERRACOTTA = "E07A5F"
SAGE = "81B29A"
SAND = "F4F1DE"
CHARCOAL = "3D405B"
MUSTARD = "F2CC8F"

FONT_JP = "游ゴシック"
FONT_EN = "Segoe UI"
FONT_MONO = "Consolas"


# ============================================================
# ヘルパー
# ============================================================
def _rgb(hex_str):
    """'E07A5F' → RGBColor"""
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def add_heading_styled(doc, text, level=1):
    """カラー付き見出しを追加"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _rgb(CHARCOAL)
        run.font.name = FONT_JP
        rPr = run._element.get_or_add_rPr()
        ea = rPr.makeelement(qn("w:rFonts"), {})
        ea.set(qn("w:eastAsia"), FONT_JP)
        rPr.append(ea)
    return h


def add_body(doc, text):
    """本文段落を追加"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = FONT_JP
        run.font.size = Pt(10.5)
    return p


def add_ascii_diagram(doc, diagram_text, font_size=Pt(8)):
    """等幅フォントの ASCII ダイアグラムを追加"""
    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = FONT_MONO
    run.font.size = font_size
    run.font.color.rgb = _rgb(CHARCOAL)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    return p


def add_colored_table(doc, headers, rows, header_color=CHARCOAL,
                      alt_color="F0F4F8", col_widths=None):
    """ヘッダー色をカスタマイズ可能なテーブル"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # ヘッダー行
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = FONT_JP
        set_cell_shading(cell, header_color)

    # データ行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = FONT_JP
            if r % 2 == 1:
                set_cell_shading(cell, alt_color)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_diagram_image(doc, png_filename, width=Inches(6.0), caption=None):
    """Excalidraw PNG ダイアグラムを挿入。ファイルがなければ False を返す"""
    png_path = os.path.join(DIAGRAMS_PNG_DIR, png_filename)
    if not os.path.isfile(png_path):
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=width)
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = _rgb(CHARCOAL)
            r.font.italic = True
    return True


def add_section_divider(doc):
    """セクション間の薄い区切り線"""
    p = doc.add_paragraph()
    run = p.add_run("_" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = _rgb(SAGE)


# ============================================================
# セクション生成関数
# ============================================================

def section_cover(doc):
    """表紙"""
    add_cover_page(
        doc,
        doc_id="TSDT-ARCH-001",
        doc_title_text="TORASAN アーキテクチャ仕様書",
        version=FRAMEWORK_VERSION,
        extra_info=[
            ("分類", "設計文書"),
            ("対象", "TORASAN 機能安全開発フレームワーク"),
        ],
    )


def section_01_overview(doc):
    """1. 概要"""
    add_heading_styled(doc, "1. 概要", level=1)

    add_body(doc,
        "TORASAN は ISO 26262 + Automotive SPICE をプロセス基盤とし、"
        "AI（Claude Code）が全フェーズを自動実行する機能安全開発フレームワークである。"
        "試験題材として洗濯機 BLDC モータ制御（IEC 60730 Class B）を扱う。"
    )

    add_heading_styled(doc, "1.1 3層アーキテクチャ", level=2)
    add_body(doc,
        "TORASAN は以下の3層で構成される。プロセス層が開発活動を管理し、"
        "規格層が安全要件を定義し、実装層が製品固有の差異を吸収する。"
    )

    add_colored_table(doc,
        ["層", "役割", "主要構成要素"],
        [
            ["プロセス層 (SPICE)", "開発活動の管理・追跡",
             "Automotive SPICE v3.1, 15プロセス, 能力レベル評価"],
            ["規格層 (ISO 26262)", "安全要件・V-model 定義",
             "ISO 26262:2018, V-model 15フェーズ, FMEA/FTTI"],
            ["製品規格層 (standard)", "製品固有差異の吸収",
             "IEC 60730 Class B / IEC 61508 / MISRA C:2012"],
        ],
        header_color=CHARCOAL,
        col_widths=[4.0, 5.0, 7.0],
    )

    add_heading_styled(doc, "1.2 製品仕様", level=2)

    add_colored_table(doc,
        ["項目", "仕様"],
        [
            ["MCU", "Renesas RL78/G14 (R5F104BG) — Flash 64KB, RAM 5.5KB, 32MHz"],
            ["モータ", "BLDC 1500W, 定格 1200rpm, 過速 1500rpm, 定格 6A / 最大 8A"],
            ["製品規格", "IEC 60730 Class B（家電制御機器）"],
            ["プロセス基盤", "ISO 26262:2018 + Automotive SPICE v3.1"],
            ["コーディング規約", "MISRA C:2012"],
            ["ADC", "10bit, Vref 5V, 電流100Hz / 電圧10Hz"],
            ["WDT", "タイムアウト 100ms, キック間隔 50ms"],
            ["RAM テスト", "March-C, 256B ブロック x22, 全体 2.2 秒"],
        ],
        header_color=TERRACOTTA,
        col_widths=[4.0, 12.0],
    )

    doc.add_page_break()


def section_02_system_architecture(doc):
    """2. システム全体構成"""
    add_heading_styled(doc, "2. システム全体構成", level=1)

    add_body(doc,
        "開発環境は Windows 11 ホスト上で動作し、WSL Ubuntu と MCP ブリッジで連携する。"
        "以下に主要コンポーネントの関係を示す。"
    )

    add_heading_styled(doc, "2.1 コンポーネント構成図", level=2)

    if not add_diagram_image(doc, "01_repo_structure.png",
                              caption="図 2-1: TORASAN リポジトリ構成（Excalidraw）"):
        # PNG がない場合は ASCII フォールバック
        diagram = (
            f"  TORASAN/  →  .claude/skills/ ({SKILL_COUNT_TOTAL}) | "
            f".claude/knowledge/ ({KNOWLEDGE_COUNT_TOTAL}) | app/ | scripts/ | docs/"
        )
        add_ascii_diagram(doc, diagram, font_size=Pt(7.5))

    add_heading_styled(doc, "2.2 リポジトリ構成", level=2)

    tree = (
        "TORASAN/\n"
        "├── CLAUDE.md              # プロジェクト指示書 (v5.0)\n"
        "├── PROCESS.md             # ISO 26262 + SPICE プロセス定義\n"
        "├── project.json           # Single Source of Truth\n"
        "├── install.sh             # 汎用スキル/ナレッジ配布\n"
        "├── .claude/\n"
        f"│   ├── skills/            # 全{SKILL_COUNT_TOTAL}スキル ({SKILL_COUNT_UNIVERSAL}汎用 + {SKILL_COUNT_DOMAIN}ドメイン)\n"
        f"│   ├── knowledge/         # 全{KNOWLEDGE_COUNT_TOTAL}ナレッジ ({KNOWLEDGE_COUNT_UNIVERSAL}汎用 + {KNOWLEDGE_COUNT_DOMAIN}ドメイン)\n"
        "│   ├── rules/             # Claude Code ルール\n"
        "│   ├── settings.json      # 環境設定\n"
        "│   └── launch.json        # Preview サーバー設定\n"
        "├── app/\n"
        "│   ├── src/               # React フロントエンド\n"
        "│   ├── server/            # Fastify API サーバー\n"
        "│   └── package.json\n"
        "├── scripts/\n"
        "│   ├── wsl-mcp-server/    # WSL エージェント MCP Server\n"
        "│   ├── generate_pptx.py   # PPTX 生成\n"
        "│   └── docx_utils.py      # Word 文書ユーティリティ\n"
        "├── docs/                  # 成果物文書\n"
        "└── process_records/       # SPICE プロセス記録"
    )
    add_ascii_diagram(doc, tree, font_size=Pt(8))

    doc.add_page_break()


def section_03_skill_system(doc):
    """3. スキルシステム"""
    add_heading_styled(doc, "3. スキルシステム", level=1)

    add_body(doc,
        f"TORASAN のスキルシステムは {SKILL_COUNT_TOTAL} 本のスキルで構成される。"
        f"汎用スキル {SKILL_COUNT_UNIVERSAL} 本は install.sh で全プロジェクトに配布され、"
        f"ドメインスキル {SKILL_COUNT_DOMAIN} 本は /repo-manage sync で個別プロジェクトに配布される。"
    )

    add_heading_styled(doc, f"3.1 汎用スキル（{SKILL_COUNT_UNIVERSAL}本）", level=2)
    add_body(doc, "全プロジェクト共通の運用基盤スキル。~/.claude/skills/ に配布される。")

    universal_skills = [
        ["session", "セッション管理", "セッション開始/終了/反省会/引継ぎ"],
        ["dashboard", "ダッシュボード", "プロジェクト概要表示"],
        ["health-check", "健全性監査", "プロジェクト健全性の包括的チェック"],
        ["skill-manage", "スキル管理", "スキル品質監査・一覧表示"],
        ["skill-evolve", "スキル改善", "PDCA 改善サイクル実行"],
        ["repo-manage", "リポジトリ管理", "PJ レジストリ・スキル同期・ドメインスキル配布"],
        ["backup", "バックアップ", "Git タグによるスナップショット"],
        ["commit-change", "コミット", "コミットワークフロー"],
        ["config-audit", "設定監査", "設定ファイル監査"],
        ["worktree-cleanup", "WT 整理", "ワークツリー整理・ブランチ削除"],
        ["update-record", "記録更新", "プロセス記録更新"],
        ["env-check", "環境チェック", "開発ツールバージョン検証"],
        ["platform-info", "プラットフォーム", "OS/シェル情報表示"],
        ["claude-master", "技術巡回", "公式ドキュメント巡回・最新化"],
    ]

    add_colored_table(doc,
        ["スキル名", "カテゴリ", "説明"],
        universal_skills,
        header_color=SAGE,
        col_widths=[3.5, 3.5, 9.0],
    )

    add_heading_styled(doc, f"3.2 ドメインスキル（{SKILL_COUNT_DOMAIN}本）", level=2)
    add_body(doc, "機能安全開発専門のスキル。.claude/skills/ に格納され /repo-manage sync で個別配布。")

    domain_skills = [
        # プロセス管理
        ["execute-phase", "プロセス管理", "V-model フェーズ実行"],
        ["assess-spice", "プロセス管理", "SPICE 能力レベル評価"],
        ["manage-tbd", "プロセス管理", "TBD 項目管理・解決追跡"],
        ["validate", "プロセス管理", "妥当性確認"],
        ["problem-resolve", "プロセス管理", "問題解決プロセス"],
        # 安全要件
        ["select-standard", "安全要件", "適用規格選定"],
        ["safety-concept", "安全要件", "安全コンセプト策定"],
        ["srs-generate", "安全要件", "SW 安全要求仕様書生成"],
        ["fmea", "安全要件", "故障モード影響解析"],
        ["switch-standard", "安全要件", "規格パッケージ切替"],
        # 設計
        ["system-design", "設計", "システム設計"],
        ["sw-design", "設計", "SW アーキテクチャ設計"],
        ["motor-control", "設計", "モータ制御アルゴリズム"],
        ["mcu-config", "設計", "MCU コンフィグレーション"],
        ["hw-review", "設計", "HW 設計レビュー"],
        # テスト・検証
        ["test-design", "テスト・検証", "テスト設計"],
        ["systest-design", "テスト・検証", "システムテスト設計"],
        ["test-coverage", "テスト・検証", "テストカバレッジ分析"],
        ["safety-verify", "テスト・検証", "機能安全検証"],
        # コード・分析
        ["driver-gen", "コード・分析", "ドライバコード生成"],
        ["static-analysis", "コード・分析", "静的解析（MISRA C 準拠）"],
        ["safety-diag", "コード・分析", "安全診断メカニズム設計"],
        # 文書・その他
        ["trace", "文書・その他", "トレーサビリティ管理"],
        ["memory-map", "文書・その他", "メモリマップ設計"],
        ["generate-docs", "文書・その他", "成果物文書生成"],
        ["ingest", "文書・その他", "外部文書取込・ナレッジ化"],
    ]

    add_colored_table(doc,
        ["スキル名", "カテゴリ", "説明"],
        domain_skills,
        header_color=TERRACOTTA,
        col_widths=[3.5, 3.0, 9.5],
    )

    add_heading_styled(doc, "3.3 配布フロー", level=2)
    add_body(doc,
        "汎用スキル・ナレッジは install.sh により ~/.claude/ にコピーされ、"
        "全プロジェクトで自動利用可能となる。ドメインスキルは /repo-manage sync コマンドで"
        "対象プロジェクトの .claude/skills/ に選択的に配布される。"
    )

    if not add_diagram_image(doc, "02_distribution_model.png",
                              caption="図 3-1: 二層配布モデル（Excalidraw）"):
        diagram = (
            f"  TORASAN/.claude/skills/ ({SKILL_COUNT_TOTAL}本)\n"
            "           │\n"
            "    ┌──────┴──────┐\n"
            "    │             │\n"
            "    ▼             ▼\n"
            "install.sh    /repo-manage sync\n"
            "    │             │\n"
            "    ▼             ▼\n"
            "~/.claude/    PJ-A/.claude/skills/\n"
            f" skills/({SKILL_COUNT_UNIVERSAL})  PJ-B/.claude/skills/\n"
            f" knowledge/({KNOWLEDGE_COUNT_UNIVERSAL}) PJ-C/.claude/skills/"
        )
        add_ascii_diagram(doc, diagram)

    doc.add_page_break()


def section_04_gui_app(doc):
    """4. 管理 GUI アプリ"""
    add_heading_styled(doc, "4. 管理 GUI アプリ", level=1)

    add_body(doc,
        "管理 GUI は読み取り専用のダッシュボードとして機能し、"
        "project.json を Single Source of Truth として全情報を可視化する。"
        "PC 起動時に Startup スクリプトで自動起動される。"
    )

    add_heading_styled(doc, "4.1 技術スタック", level=2)

    add_colored_table(doc,
        ["技術", "バージョン", "用途"],
        [
            ["React", "19.2", "UI フレームワーク"],
            ["Vite", "7.3", "ビルドツール / 開発サーバー (:3000)"],
            ["Fastify", "5.3", "API サーバー (:3001)"],
            ["Redux Toolkit", "2.11", "状態管理"],
            ["Tailwind CSS", "4.2", "スタイリング"],
            ["docx", "9.6", "Word 文書エクスポート"],
            ["ExcelJS", "4.4", "Excel エクスポート"],
            ["chokidar", "5.0", "ファイル監視"],
            ["TypeScript", "5.9", "型安全"],
            ["react-router-dom", "7.13", "ルーティング"],
        ],
        header_color=CHARCOAL,
        col_widths=[3.5, 2.5, 10.0],
    )

    add_heading_styled(doc, "4.2 ページ構成（12ページ）", level=2)

    add_colored_table(doc,
        ["#", "ページ", "パス", "説明"],
        [
            ["1", "ダッシュボード", "/", "プロジェクト概要・進捗サマリ"],
            ["2", "フェーズ", "/phases", "V-model 15 フェーズ進捗管理"],
            ["3", "要件", "/requirements", "要件エディタ・CRUD"],
            ["4", "トレーサビリティ", "/trace", "要件-設計-テスト追跡マトリクス"],
            ["5", "文書", "/documents", "成果物文書生成・エクスポート"],
            ["6", "レビュー", "/reviews", "レビュー記録"],
            ["7", "変更履歴", "/changelog", "changeLog 表示"],
            ["8", "スキル", "/skills", "スキル一覧・成熟度表示"],
            ["9", "セッション", "/sessions", "セッション履歴"],
            ["10", "トラッカー", "/tracker", "Ideas / Issues 管理"],
            ["11", "レジストリ", "/registry", "プロジェクトレジストリ"],
            ["12", "環境管理", "/environment", "開発ツールバージョン表示"],
        ],
        header_color=SAGE,
        col_widths=[1.0, 3.0, 3.5, 8.5],
    )

    add_heading_styled(doc, "4.3 API エンドポイント", level=2)

    add_colored_table(doc,
        ["メソッド", "エンドポイント", "説明"],
        [
            ["GET", "/api/project", "project.json 全体取得"],
            ["PATCH", "/api/project", "project.json フィールド更新"],
            ["GET", "/api/changelog", "変更履歴取得"],
            ["GET", "/api/skills", "スキル一覧取得"],
            ["GET", "/api/sessions", "セッション履歴取得"],
            ["GET", "/api/registry", "プロジェクトレジストリ取得"],
            ["GET", "/api/environment", "開発環境情報取得"],
            ["GET", "/api/tracker", "Ideas / Issues 取得"],
            ["WS", "/ws", "WebSocket リアルタイム更新"],
        ],
        header_color=CHARCOAL,
        col_widths=[2.5, 4.0, 9.5],
    )

    add_heading_styled(doc, "4.4 リアルタイム更新パイプライン", level=2)
    add_body(doc,
        "project.json の変更は chokidar で検知され、WebSocket 経由で"
        "ブラウザに1秒以内にプッシュされる。"
    )

    pipeline = (
        "project.json\n"
        "     │\n"
        "     │ chokidar (ファイル監視)\n"
        "     ▼\n"
        "Fastify Server (:3001)\n"
        "     │\n"
        "     │ WebSocket push\n"
        "     ▼\n"
        "Redux Store (RTK)\n"
        "     │\n"
        "     │ useSelector\n"
        "     ▼\n"
        "React Components (再レンダリング)"
    )
    add_ascii_diagram(doc, pipeline)

    doc.add_page_break()


def section_05_wsl_agent(doc):
    """5. WSL エージェント連携"""
    add_heading_styled(doc, "5. WSL エージェント連携", level=1)

    add_body(doc,
        "Windows 上の Claude Code から WSL Ubuntu 上の Claude Code に"
        "タスクを委任するための MCP Server ベースの連携システム。"
        "stdio トランスポートで通信し、非同期タスク実行をサポートする。"
    )

    add_heading_styled(doc, "5.1 アーキテクチャ", level=2)

    diagram = (
        "┌───────────────────────┐\n"
        "│  Claude Code          │\n"
        "│  (Windows Host)       │\n"
        "│                       │\n"
        "│  MCP Client           │\n"
        "└───────────┬───────────┘\n"
        "            │ stdio (stdin/stdout)\n"
        "┌───────────▼───────────┐\n"
        "│  MCP Server           │\n"
        "│  wsl-mcp-server/      │\n"
        "│  Node.js + TypeScript │\n"
        "│                       │\n"
        "│  ┌─────────────────┐  │\n"
        "│  │  Task Queue     │  │\n"
        "│  │  (In-Memory)    │  │\n"
        "│  └────────┬────────┘  │\n"
        "│           │           │\n"
        "│  ┌────────▼────────┐  │\n"
        "│  │  WSL Executor   │  │\n"
        "│  │  wsl.exe bridge │  │\n"
        "│  └────────┬────────┘  │\n"
        "└───────────┼───────────┘\n"
        "            │ wsl.exe\n"
        "┌───────────▼───────────┐\n"
        "│  WSL Ubuntu 24.04     │\n"
        "│                       │\n"
        "│  Claude Code Agent    │\n"
        "│  (claude --print)     │\n"
        "│                       │\n"
        "│  OAuth Token 認証     │\n"
        "│  ~/.claude/oauth_token│\n"
        "└───────────────────────┘"
    )
    add_ascii_diagram(doc, diagram, font_size=Pt(7.5))

    add_heading_styled(doc, "5.2 MCP ツール一覧", level=2)

    add_colored_table(doc,
        ["ツール", "説明", "パラメータ"],
        [
            ["delegate_task",
             "WSL エージェントにタスクを委任",
             "task, cwd, model, max_turns, timeout_seconds, async"],
            ["check_task_status",
             "非同期タスクの状態確認",
             "task_id"],
            ["get_task_result",
             "完了タスクの結果取得",
             "task_id"],
            ["list_tasks",
             "全タスク一覧",
             "(なし)"],
        ],
        header_color=CHARCOAL,
        col_widths=[3.5, 5.0, 7.5],
    )

    add_heading_styled(doc, "5.3 タスク状態遷移", level=2)

    state_diagram = (
        "  queued ──────▶ running ──────▶ completed\n"
        "                   │\n"
        "                   ├──────────▶ failed\n"
        "                   │\n"
        "                   └──────────▶ timeout"
    )
    add_ascii_diagram(doc, state_diagram)

    add_body(doc,
        "async=true の場合、delegate_task は即座に task_id を返し、"
        "バックグラウンドで WSL エージェントを実行する。"
        "check_task_status で状態を確認し、completed 後に get_task_result で結果を取得する。"
    )

    doc.add_page_break()


def section_06_data_model(doc):
    """6. データモデル"""
    add_heading_styled(doc, "6. データモデル", level=1)

    add_body(doc,
        "project.json はプロジェクトの Single Source of Truth であり、"
        "全構成情報・進捗・トレーサビリティ・変更履歴を一元管理する。"
        "データは 3 つのスコープ（Session / Project / Global）に分類される。"
    )

    add_diagram_image(doc, "03_memory_architecture.png", width=Inches(5.5),
                      caption="図 6-1: メモリ・状態管理アーキテクチャ（Excalidraw）")

    add_heading_styled(doc, "6.1 トップレベルキー", level=2)

    add_colored_table(doc,
        ["キー", "型", "説明"],
        [
            ["_schema_version", "string", "スキーマバージョン (現在 3.0)"],
            ["projectName", "string", "プロジェクト名称"],
            ["category", "string", "プロジェクト分類"],
            ["standard", "object", "適用規格設定 (base, product_override, spice_level_target)"],
            ["standard", "object", "適用規格設定 (base, product_override, product_override_class)"],
            ["product_spec", "object", "製品仕様 (mcu, motor, adc, wdt)"],
            ["traceability", "object", "要件追跡カウント (sg, fsr, tsr, sr, tc)"],
            ["tbd_count / tbd_items", "number / array", "未確定項目数・一覧"],
            ["ideas", "array<IdeaJson>", "アイデア一覧"],
            ["issues", "array<IssueJson>", "課題一覧"],
            ["phases", "object<PhaseEntry>", "15フェーズ進捗 (PH-01 ~ PH-15)"],
            ["changeLog", "array", "変更履歴"],
            ["spice_assessment", "object", "SPICE 能力レベル評価結果"],
            ["gateLog", "array", "ゲートレビュー記録"],
            ["problemLog", "array", "問題記録"],
            ["reviewLog", "array", "レビュー記録"],
        ],
        header_color=CHARCOAL,
        col_widths=[4.0, 3.5, 8.5],
    )

    add_heading_styled(doc, "6.2 PhaseEntry フィールド", level=2)

    add_colored_table(doc,
        ["フィールド", "型", "説明", "例"],
        [
            ["name", "string", "フェーズ名称", "プロジェクト計画・安全計画"],
            ["spice", "string", "SPICE プロセス ID", "MAN.3"],
            ["status", "string", "進捗状態", "completed / not_started / in_progress"],
            ["docStatus", "string", "文書版数", "v2.1 / --"],
            ["gateStatus", "string", "ゲート結果", "PASS / FAIL / --"],
        ],
        header_color=SAGE,
        col_widths=[3.0, 2.0, 5.5, 5.5],
    )

    add_heading_styled(doc, "6.3 IdeaJson / IssueJson フィールド", level=2)

    add_colored_table(doc,
        ["フィールド", "IdeaJson", "IssueJson"],
        [
            ["id", "IDEA-001 ...", "ISS-001 ..."],
            ["title", "アイデアタイトル", "課題タイトル"],
            ["description", "詳細説明", "詳細説明"],
            ["status", "open / dismissed", "open / resolved"],
            ["source", "retrospective / version-audit", "retrospective"],
            ["severity", "(なし)", "low / medium / high"],
            ["sessionDate", "セッション日付", "セッション日付"],
            ["resolvedDate", "(なし)", "解決日"],
            ["resolution", "(なし)", "解決内容"],
            ["dismissReason", "却下理由", "(なし)"],
        ],
        header_color=MUSTARD,
        col_widths=[3.5, 6.0, 6.5],
    )

    doc.add_page_break()


def section_07_vmodel(doc):
    """7. V字モデル（15フェーズ）"""
    add_heading_styled(doc, "7. V字モデル（15フェーズ）", level=1)

    add_body(doc,
        "ISO 26262 の V-model に基づく 15 フェーズで構成される。"
        "各フェーズは SPICE プロセスに対応し、ゲートレビューを通過して完了となる。"
    )

    phases = [
        ["PH-01", "プロジェクト計画・安全計画", "MAN.3", "completed", "v2.1", "PASS"],
        ["PH-02", "アイテム定義", "SYS.1", "completed", "v2.0", "PASS"],
        ["PH-03", "HARA", "SYS.1", "completed", "v2.1", "PASS"],
        ["PH-04", "FSC", "SYS.2", "completed", "v2.1", "PASS"],
        ["PH-05", "TSC", "SYS.3", "not_started", "\u2014", "\u2014"],
        ["PH-06", "システム設計", "SYS.3", "not_started", "\u2014", "\u2014"],
        ["PH-07", "HW設計", "SYS.3", "not_started", "\u2014", "\u2014"],
        ["PH-08", "SW安全要求(SRS)", "SWE.1", "not_started", "\u2014", "\u2014"],
        ["PH-09", "SWアーキテクチャ設計", "SWE.2", "not_started", "\u2014", "\u2014"],
        ["PH-10", "SWユニット設計・実装", "SWE.3", "not_started", "\u2014", "\u2014"],
        ["PH-11", "SWテスト仕様・実施", "SWE.4-6", "not_started", "\u2014", "\u2014"],
        ["PH-12", "HW統合テスト", "SYS.4", "not_started", "\u2014", "\u2014"],
        ["PH-13", "システムテスト", "SYS.5", "not_started", "\u2014", "\u2014"],
        ["PH-14", "機能安全検証", "SYS.5", "not_started", "\u2014", "\u2014"],
        ["PH-15", "安全アセスメント", "SUP.1", "not_started", "\u2014", "\u2014"],
    ]

    add_colored_table(doc,
        ["ID", "フェーズ名", "SPICE", "Status", "Doc", "Gate"],
        phases,
        header_color=CHARCOAL,
        col_widths=[1.8, 5.0, 2.0, 2.5, 1.5, 1.5],
    )

    add_note(doc, "PH-01 ~ PH-04 は 2 周目レビュー完了済み（2026-03-02）。PH-05 以降は未着手。")

    add_heading_styled(doc, "7.1 V字モデル構成", level=2)

    if not add_diagram_image(doc, "04_vmodel.png", width=Inches(5.0),
                              caption="図 7-1: SPICE V字モデル（Excalidraw）"):
        v_diagram = (
            "  PH-01  計画                                          PH-15 安全ｱｾｽ\n"
            "    \\                                                    /\n"
            "  PH-02  ｱｲﾃﾑ定義                              PH-14 安全検証\n"
            "      \\                                            /\n"
            "    PH-03  HARA                            PH-13 ｼｽﾃﾑﾃｽﾄ\n"
            "        \\                                      /\n"
            "      PH-04  FSC                      PH-12 HW統合ﾃｽﾄ\n"
            "          \\                                /\n"
            "        PH-05  TSC                PH-11 SWﾃｽﾄ\n"
            "            \\                          /\n"
            "          PH-06  ｼｽﾃﾑ設計      PH-10 SWﾕﾆｯﾄ\n"
            "              \\                    /\n"
            "            PH-07  HW設計  PH-09 SWｱｰｷ\n"
            "                \\            /\n"
            "              PH-08  SRS（底部）"
        )
        add_ascii_diagram(doc, v_diagram, font_size=Pt(7.5))

    doc.add_page_break()


def section_08_knowledge(doc):
    """8. ナレッジベース"""
    add_heading_styled(doc, "8. ナレッジベース", level=1)

    add_body(doc,
        "ナレッジはスキル実行時に参照されるリファレンス情報である。"
        "2 階層構造で管理され、汎用ナレッジは全 PJ で、ドメインナレッジは TORASAN 管理下で利用される。"
    )

    add_heading_styled(doc, f"8.1 汎用ナレッジ（{KNOWLEDGE_COUNT_UNIVERSAL}本）", level=2)
    add_body(doc, "~/.claude/knowledge/ に配布。全プロジェクトで自動参照される。")

    add_colored_table(doc,
        ["#", "ファイル名", "内容"],
        [
            ["1", "claude_code_ops.md", "コンテキスト管理・メモリ設計・トークン最適化"],
            ["2", "claude_platform_updates.md", "技術巡回履歴・プラットフォーム更新情報"],
            ["3", "skill_lifecycle.md", "スキル成熟度モデル (L1-L5) 定義"],
            ["4", "skill_feedback_log.md", "スキル実行フィードバック記録"],
            ["5", "git_worktree_branch_management.md", "Git ワークツリー・ブランチ戦略"],
            ["6", "cross_platform_dev.md", "Windows/WSL/macOS 互換性ガイド"],
            ["7", "memory_paths.md", "メモリパス規約・slug 算出ルール"],
            ["8", "error_prevention.md", "エラー防止・PDCA プロセス定義"],
        ],
        header_color=SAGE,
        col_widths=[1.0, 6.5, 8.5],
    )

    add_heading_styled(doc, f"8.2 ドメインナレッジ（{KNOWLEDGE_COUNT_DOMAIN}本）", level=2)
    add_body(doc, ".claude/knowledge/ に格納。機能安全専門の規格・技術リファレンス。")

    add_colored_table(doc,
        ["#", "ファイル名", "内容"],
        [
            ["1", "iso26262_iec60730.md", "ISO 26262 / IEC 60730 規格要件"],
            ["2", "product_standard_mapping.md", "製品カテゴリ→規格マッピング"],
            ["3", "automotive_spice.md", "Automotive SPICE プロセスモデル"],
            ["4", "safety_case_gsn.md", "GSN 安全論証ノーテーション"],
            ["5", "safety_diagnostics.md", "安全診断メカニズム設計"],
            ["6", "misra_c_2012.md", "MISRA C:2012 コーディング規約"],
            ["7", "fmea_guide.md", "FMEA 故障モード影響解析ガイド"],
            ["8", "bldc_safety.md", "BLDC モータ制御安全要件"],
            ["9", "srs_template.md", "SRS テンプレート・要件仕様"],
            ["10", "pptx_advanced_shapes.md", "python-pptx 高度図形操作リファレンス"],
            ["11", "uml_diagramming.md", "UML 作図リファレンス（python-pptx 実装パターン）"],
            ["12", "excalidraw_pipeline.md", "Excalidraw → DOCX ダイアグラムパイプライン"],
            ["13", "md2pptx_guide.md", "Markdown → PPTX 変換ガイド"],
        ],
        header_color=TERRACOTTA,
        col_widths=[1.0, 6.5, 8.5],
    )

    add_heading_styled(doc, "8.3 ナレッジ参照の仕組み", level=2)
    add_body(doc,
        "スキルは Markdown 内の参照指定に基づき、実行時に必要なナレッジを自動的に読み込む。"
        "例えば /driver-gen は misra_c_2012.md と bldc_safety.md を参照する。"
    )

    doc.add_page_break()


def section_09_startup_notify(doc):
    """9. 起動自動化・通知"""
    add_heading_styled(doc, "9. 起動自動化・通知", level=1)

    add_heading_styled(doc, "9.1 PC 起動時自動起動", level=2)
    add_body(doc,
        "Windows の Startup フォルダに VBS スクリプトを配置し、"
        "PC 起動時に管理 GUI アプリと通知監視デーモンを自動起動する。"
    )

    add_procedure_box(doc, "起動フロー", [
        "Windows 起動 → Startup フォルダの VBS 実行",
        "start_torasan_app.vbs → npm run dev (Vite :3000 + Fastify :3001)",
        "ブラウザで localhost:3000 を自動オープン",
        "start_notify_monitor.vbs → Python 通知デーモン起動",
    ])

    add_heading_styled(doc, "9.2 通知デーモン", level=2)
    add_body(doc,
        "Claude Code の停止フック（settings.json の Stop イベント）は"
        "本体のバグにより動作しない（Issue #29560, #26770）。"
        "代替として Python デーモンがデバッグログを監視し、"
        "タスク完了時に通知音を再生する。"
    )

    add_colored_table(doc,
        ["コンポーネント", "パス", "役割"],
        [
            ["通知デーモン", "~/.claude/claude_notify_monitor.py", "デバッグログ監視 → 通知音再生"],
            ["起動スクリプト", "~/.claude/claude-notify.sh", "start / stop / status / log"],
            ["VBS ラッパー", "Startup/start_notify_monitor.vbs", "Windows 起動時の自動起動"],
        ],
        header_color=CHARCOAL,
        col_widths=[3.5, 7.0, 5.5],
    )

    diagram = (
        "Claude Code (実行完了)\n"
        "        │\n"
        "        ▼\n"
        "  debug.log (追記)\n"
        "        │\n"
        "        │  claude_notify_monitor.py (tail -F 相当)\n"
        "        ▼\n"
        "  完了パターン検出\n"
        "        │\n"
        "        ▼\n"
        "  通知音再生 (winsound / afplay)"
    )
    add_ascii_diagram(doc, diagram)

    doc.add_page_break()


def section_10_tech_stack(doc):
    """10. 技術スタック"""
    add_heading_styled(doc, "10. 技術スタック", level=1)

    add_heading_styled(doc, "10.1 ツールパス解決フロー", level=2)
    add_body(doc,
        "Windows 環境ではツールの実行パスが複雑になるため、"
        "以下のフローで自動解決を行う。"
    )
    add_diagram_image(doc, "05_tool_resolution.png", width=Inches(4.0),
                      caption="図 10-1: ツールパス解決フロー（Excalidraw）")

    add_heading_styled(doc, "10.2 開発環境ツール", level=2)

    add_colored_table(doc,
        ["ツール", "バージョン", "用途"],
        [
            ["Node.js", "24.14.0", "JavaScript ランタイム"],
            ["npm", "11.9.0", "パッケージマネージャ"],
            ["Python", "3.14.3", "スクリプト・ツール"],
            ["TypeScript", "5.9.3", "型安全な JavaScript"],
            ["Git", "2.53.0", "バージョン管理"],
            ["cppcheck", "2.20.0", "C/C++ 静的解析"],
            ["clang-tidy", "21.1.8", "C/C++ リンター"],
            ["flawfinder", "2.0.19", "セキュリティ脆弱性スキャン"],
            ["gh (GitHub CLI)", "2.87.3", "GitHub 操作"],
            ["MSYS2", "\u2014", "Windows Unix ツール"],
            ["WSL", "Ubuntu 24.04", "Linux サブシステム"],
        ],
        header_color=CHARCOAL,
        col_widths=[3.5, 3.0, 9.5],
    )

    add_heading_styled(doc, "10.3 app/ npm パッケージ", level=2)

    add_colored_table(doc,
        ["パッケージ", "バージョン", "分類"],
        [
            ["react", "^19.2.0", "dependencies"],
            ["react-dom", "^19.2.0", "dependencies"],
            ["react-redux", "^9.2.0", "dependencies"],
            ["react-router-dom", "^7.13.1", "dependencies"],
            ["@reduxjs/toolkit", "^2.11.2", "dependencies"],
            ["fastify", "^5.3.3", "dependencies"],
            ["@fastify/websocket", "^11.2.0", "dependencies"],
            ["tailwindcss", "^4.2.1", "dependencies"],
            ["@tailwindcss/vite", "^4.2.1", "dependencies"],
            ["chokidar", "^5.0.0", "dependencies"],
            ["docx", "^9.6.0", "dependencies"],
            ["exceljs", "^4.4.0", "dependencies"],
            ["gray-matter", "^4.0.3", "dependencies"],
            ["vite", "^7.3.1", "devDependencies"],
            ["typescript", "~5.9.3", "devDependencies"],
            ["eslint", "^9.39.1", "devDependencies"],
            ["tsx", "^4.21.0", "devDependencies"],
            ["concurrently", "^9.2.1", "devDependencies"],
        ],
        header_color=SAGE,
        col_widths=[4.5, 2.5, 3.5],
    )

    add_heading_styled(doc, "10.4 wsl-mcp-server/ npm パッケージ", level=2)

    add_colored_table(doc,
        ["パッケージ", "バージョン", "分類"],
        [
            ["@modelcontextprotocol/sdk", "^1.12.0", "dependencies"],
            ["zod", "^3.24.0", "dependencies"],
            ["tsx", "^4.21.0", "devDependencies"],
            ["typescript", "~5.9.3", "devDependencies"],
            ["@types/node", "^24.0.0", "devDependencies"],
        ],
        header_color=TERRACOTTA,
        col_widths=[5.5, 2.5, 3.5],
    )


# ============================================================
# メイン
# ============================================================
def main():
    doc = Document()
    setup_document(doc)

    # 表紙
    section_cover(doc)

    # 本文セクション
    section_01_overview(doc)
    section_02_system_architecture(doc)
    section_03_skill_system(doc)
    section_04_gui_app(doc)
    section_05_wsl_agent(doc)
    section_06_data_model(doc)
    section_07_vmodel(doc)
    section_08_knowledge(doc)
    section_09_startup_notify(doc)
    section_10_tech_stack(doc)

    # 保存
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "TORASAN_Architecture.docx")
    doc.save(output_path)
    print(f"Generated: {output_path}")
    print(f"Sections: 10 (+ cover page)")


if __name__ == "__main__":
    main()
