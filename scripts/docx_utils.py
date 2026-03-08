"""Word文書生成用の共通ユーティリティ"""

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """ヘッダー付きスタイルテーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2E4057')

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_procedure_box(doc, title, steps):
    """手順ボックス（タイトル + ステップリスト）"""
    p = doc.add_paragraph()
    run = p.add_run(f'  {title}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'  Step {i}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(step)
        run.font.size = Pt(10)


def add_command_block(doc, command, description=''):
    """コマンドブロック（等幅フォント + 緑色）"""
    if description:
        doc.add_paragraph(description)
    p = doc.add_paragraph()
    run = p.add_run(f'    {command}')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x44)


def add_note(doc, text):
    """NOTE ブロック"""
    p = doc.add_paragraph()
    run = p.add_run('NOTE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run.font.size = Pt(9.5)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def setup_document(doc):
    """文書の共通初期設定（ページ設定・フォント）"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Yu Gothic'
        h_style.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)


def add_cover_page(doc, doc_id, doc_title_text, version, extra_info=None):
    """表紙を追加"""
    from datetime import date as _date

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TORASAN')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('機能安全開発プロセスフレームワーク')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    doc.add_paragraph()

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dt.add_run(doc_title_text)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for _ in range(4):
        doc.add_paragraph()

    info_rows = [
        ('文書番号', doc_id),
        ('版数', version),
        ('作成日', _date.today().strftime('%Y-%m-%d')),
    ]
    if extra_info:
        info_rows.extend(extra_info)

    info_table = doc.add_table(rows=len(info_rows), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'
    for i, (k, v) in enumerate(info_rows):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        set_cell_shading(info_table.rows[i].cells[0], 'E8EDF2')
        for cell in info_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_page_break()
