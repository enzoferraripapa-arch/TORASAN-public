#!/usr/bin/env python3
"""
成果物エクスポートスクリプト
Markdown → Word (.docx) / Excel (.xlsx) 変換 + SPICEプロセスフォルダ構成
"""
import os
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# ─── 設定 ───────────────────────────────────────────
PROJECT_NAME = "洗濯機BLDCモータ制御"
BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "docs"
OUTPUT_BASE = Path(r"D:\成果物") / PROJECT_NAME

# SPICEプロセスフォルダマッピング（設計成果物: docs/ から取得）
SPICE_FOLDERS = {
    "MAN.3_プロジェクト管理": [
        ("01_安全計画.md", "01_安全計画.docx"),
        ("ツール適格性記録.md", "ツール適格性記録.docx"),
    ],
    "SYS.1_要件抽出": [
        ("02_アイテム定義.md", "02_アイテム定義.docx"),
        ("03_HARA.md", "03_HARA.docx"),
    ],
    "SYS.2_システム要件分析": [
        ("04_FSC.md", "04_FSC.docx"),
        ("04a_システム要件仕様.md", "04a_システム要件仕様.docx"),
    ],
    "SYS.3_システムアーキテクチャ設計": [
        ("05_TSC.md", "05_TSC.docx"),
        ("06_システム設計.md", "06_システム設計.docx"),
        ("07_HW設計.md", "07_HW設計.docx"),
    ],
    "SWE.1_ソフトウェア要件分析": [
        ("08_SW安全要求.md", "08_SW安全要求仕様書.docx"),
    ],
    "SWE.2_ソフトウェアアーキテクチャ設計": [
        ("09_SWアーキテクチャ.md", "09_SWアーキテクチャ設計書.docx"),
    ],
    "SWE.3_ソフトウェア詳細設計": [
        ("10_SWユニット設計.md", "10_SW詳細設計書.docx"),
        ("コーディングガイドライン.md", "コーディングガイドライン.docx"),
    ],
    "SWE.4_ソフトウェアユニット検証": [
        ("11a_SWユニットテスト仕様書.md", "11a_SWユニットテスト仕様書.docx"),
    ],
    "SWE.5_ソフトウェア統合テスト": [
        ("11b_SW結合テスト仕様書.md", "11b_SW結合テスト仕様書.docx"),
    ],
    "SWE.6_ソフトウェア適格性テスト": [
        ("11c_SWコンポーネントテスト仕様書.md", "11c_SWコンポーネントテスト仕様書.docx"),
    ],
    "SYS.4_システム統合テスト": [
        ("12_HW統合テスト仕様書.md", "12_HW統合テスト仕様書.docx"),
    ],
    "SYS.5_システム適格性テスト": [
        ("13_システムテスト仕様書.md", "13_システムテスト仕様書.docx"),
    ],
    "共通_トレーサビリティ": [
        ("トレーサビリティ.md", "トレーサビリティマトリクス.xlsx"),
    ],
}

# プロセス記録マッピング（process_records/ から各SPICEフォルダへ配置）
PROCESS_RECORD_MAP = {
    "MAN.3_プロジェクト管理":           "MAN.3_project_plan.md",
    "SYS.1_要件抽出":                   "SYS.1_requirements_elicitation.md",
    "SYS.2_システム要件分析":           "SYS.2_system_requirements.md",
    "SYS.3_システムアーキテクチャ設計": "SYS.3_system_architecture.md",
    "SYS.4_システム統合テスト":         "SYS.4_system_integration.md",
    "SYS.5_システム適格性テスト":       "SYS.5_system_qualification.md",
    "SWE.1_ソフトウェア要件分析":       "SWE.1_requirements.md",
    "SWE.2_ソフトウェアアーキテクチャ設計": "SWE.2_architecture.md",
    "SWE.3_ソフトウェア詳細設計":       "SWE.3_detailed_design.md",
    "SWE.4_ソフトウェアユニット検証":   "SWE.4_unit_verification.md",
    "SWE.5_ソフトウェア統合テスト":     "SWE.5_integration_test.md",
    "SWE.6_ソフトウェア適格性テスト":   "SWE.6_qualification_test.md",
    "SUP.1_品質保証":                   "SUP.1_quality_assurance.md",
    "SUP.8_構成管理":                   "SUP.8_config_management.md",
    "SUP.9_問題解決管理":               "SUP.9_problem_resolution.md",
    "SUP.10_変更要求管理":              "SUP.10_change_request.md",
}

PR_DIR = BASE_DIR / "process_records"

# 各SPICEフォルダの説明と予定成果物（00_INDEX.md 自動生成用）
FOLDER_INDEX = {
    "MAN.3_プロジェクト管理":           {"ph": "PH-01", "desc": "プロジェクト管理プロセス（MAN.3）", "docs": ["安全計画", "ツール適格性記録"]},
    "SYS.1_要件抽出":                   {"ph": "PH-02,03", "desc": "要件抽出プロセス（SYS.1）", "docs": ["アイテム定義", "HARA"]},
    "SYS.2_システム要件分析":           {"ph": "PH-04", "desc": "システム要件分析プロセス（SYS.2）", "docs": ["機能安全コンセプト(FSC)"]},
    "SYS.3_システムアーキテクチャ設計": {"ph": "PH-05,06", "desc": "システムアーキテクチャ設計プロセス（SYS.3）", "docs": ["技術安全コンセプト(TSC)", "システム設計"]},
    "SYS.4_システム統合テスト":         {"ph": "PH-12", "desc": "システム統合テストプロセス（SYS.4）", "docs": ["HW統合テスト仕様書"]},
    "SYS.5_システム適格性テスト":       {"ph": "PH-13,14", "desc": "システム適格性テストプロセス（SYS.5）", "docs": ["システムテスト仕様書", "機能安全検証報告書(予定)"]},
    "SWE.1_ソフトウェア要件分析":       {"ph": "PH-08", "desc": "ソフトウェア要件分析プロセス（SWE.1）", "docs": ["SW安全要求仕様書(SRS)"]},
    "SWE.2_ソフトウェアアーキテクチャ設計": {"ph": "PH-09", "desc": "ソフトウェアアーキテクチャ設計プロセス（SWE.2）", "docs": ["SWアーキテクチャ設計書"]},
    "SWE.3_ソフトウェア詳細設計":       {"ph": "PH-10", "desc": "ソフトウェア詳細設計プロセス（SWE.3）", "docs": ["SW詳細設計書", "コーディングガイドライン", "ソースコード"]},
    "SWE.4_ソフトウェアユニット検証":   {"ph": "PH-11", "desc": "ソフトウェアユニット検証プロセス（SWE.4）", "docs": ["SWユニットテスト仕様書"]},
    "SWE.5_ソフトウェア統合テスト":     {"ph": "PH-11", "desc": "ソフトウェア統合テストプロセス（SWE.5）", "docs": ["SW結合テスト仕様書"]},
    "SWE.6_ソフトウェア適格性テスト":   {"ph": "PH-11", "desc": "ソフトウェア適格性テストプロセス（SWE.6）", "docs": ["SWコンポーネントテスト仕様書"]},
    "SUP.1_品質保証":                   {"ph": "PH-15", "desc": "品質保証プロセス（SUP.1）", "docs": ["安全アセスメント報告書(予定)"]},
    "SUP.8_構成管理":                   {"ph": "継続", "desc": "構成管理プロセス（SUP.8）", "docs": ["構成管理計画"]},
    "SUP.9_問題解決管理":               {"ph": "随時", "desc": "問題解決管理プロセス（SUP.9）", "docs": ["問題報告書(随時)"]},
    "SUP.10_変更要求管理":              {"ph": "随時", "desc": "変更要求管理プロセス（SUP.10）", "docs": ["変更要求書(随時)"]},
    "共通_トレーサビリティ":             {"ph": "PH-03〜", "desc": "全フェーズ横断のトレーサビリティマトリクス", "docs": ["トレーサビリティマトリクス"]},
}

# ─── Word スタイル設定 ─────────────────────────────────

def set_cell_shading(cell, color_hex):
    """テーブルセルの背景色を設定"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top="single", bottom="single", left="single", right="single", sz=4, color="auto"):
    """テーブルセルのボーダーを設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), str(sz))
        edge_el.set(qn('w:color'), color)
        edge_el.set(qn('w:space'), '0')
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def setup_document_styles(doc):
    """Word文書のスタイルを設定"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Yu Gothic'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if level == 1:
            heading_style.font.size = Pt(18)
        elif level == 2:
            heading_style.font.size = Pt(14)
        elif level == 3:
            heading_style.font.size = Pt(12)
        elif level == 4:
            heading_style.font.size = Pt(11)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)


# ─── Markdown パーサ ─────────────────────────────────

def parse_markdown_table(lines):
    """Markdownテーブルをパース"""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            # セパレータ行をスキップ
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            rows.append(cells)
    return rows


def convert_md_to_docx(md_path, docx_path):
    """MarkdownファイルをWord文書に変換"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    setup_document_styles(doc)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []

    while i < len(lines):
        line = lines[i]

        # コードブロック
        if line.strip().startswith('```'):
            if in_code_block:
                # コードブロック終了
                in_code_block = False
                code_text = '\n'.join(code_block_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    p.paragraph_format.left_indent = Cm(1.0)
                    pPr = p._element.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F5F5F5')
                    shd.set(qn('w:val'), 'clear')
                    pPr.append(shd)
                code_block_lines = []
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 水平線
        if line.strip() == '---':
            # page_break の代わりに水平線を表現
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), 'CCCCCC')
            bottom.set(qn('w:space'), '1')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 見出し
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line.strip())
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # テーブル
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            rows = parse_markdown_table(table_lines)
            if rows:
                num_cols = max(len(r) for r in rows)
                # 列数を統一
                for idx, r in enumerate(rows):
                    while len(r) < num_cols:
                        r.append('')

                table = doc.add_table(rows=len(rows), cols=num_cols, style='Table Grid')
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        # Bold処理
                        parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                p.add_run(part)

                        p.style = doc.styles['Normal']
                        for run in p.runs:
                            run.font.size = Pt(9)
                            run.font.name = 'Yu Gothic'

                        # ヘッダ行のスタイル
                        if row_idx == 0:
                            set_cell_shading(cell, '2C3E50')
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.bold = True
                        else:
                            # 偶数行の背景色
                            if row_idx % 2 == 0:
                                set_cell_shading(cell, 'F8F9FA')

                # テーブルの後にスペース
                doc.add_paragraph('')
            continue

        # 箇条書き
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)$', line.strip())
        if list_match:
            prefix = list_match.group(2)
            text = list_match.group(3)
            indent = len(re.match(r'^(\s*)', line).group(1))

            p = doc.add_paragraph()
            # Bold処理
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # インラインコード
                    code_parts = re.split(r'(`[^`]+`)', part)
                    for cp in code_parts:
                        if cp.startswith('`') and cp.endswith('`'):
                            run = p.add_run(cp[1:-1])
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                        else:
                            p.add_run(cp)

            if prefix in ['-', '*']:
                p.style = 'List Bullet'
            else:
                p.style = 'List Number'

            if indent >= 2:
                p.paragraph_format.left_indent = Cm(1.5)

            i += 1
            continue

        # 通常テキスト
        p = doc.add_paragraph()
        text = line.strip()
        # Bold + インラインコード処理
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            else:
                p.add_run(part)

        i += 1

    doc.save(str(docx_path))
    print(f"  [OK] {docx_path.name}")


def convert_trace_to_xlsx(md_path, xlsx_path):
    """トレーサビリティマトリクスをExcelに変換"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    wb = openpyxl.Workbook()

    # スタイル定義
    header_font = Font(name='Yu Gothic', bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
    normal_font = Font(name='Yu Gothic', size=10)
    title_font = Font(name='Yu Gothic', bold=True, size=14, color='1A1A2E')
    subtitle_font = Font(name='Yu Gothic', bold=True, size=12, color='2C3E50')
    pass_fill = PatternFill(start_color='D4EDDA', end_color='D4EDDA', fill_type='solid')
    tbd_fill = PatternFill(start_color='FFF3CD', end_color='FFF3CD', fill_type='solid')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin'),
    )
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    wrap_align = Alignment(wrap_text=True, vertical='top')

    # セクションをパースしてシートに分割
    sections = re.split(r'\n## ', content)

    # --- シート1: HZ→SG ---
    ws1 = wb.active
    ws1.title = "HZ→SG"
    ws1['A1'] = 'HZ → SG トレーサビリティ（PH-03 確立）'
    ws1['A1'].font = title_font
    ws1.merge_cells('A1:F1')

    # 順方向テーブル
    for section in sections:
        if '順方向トレース（ハザード → 安全目標）' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = 3
                ws1.cell(row=start_row - 1, column=1, value='順方向トレース（HZ → SG）').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws1.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font
                            if '**PASS**' in val or 'OK' in val:
                                cell.fill = pass_fill
                            elif 'TBD' in val:
                                cell.fill = tbd_fill

        if '逆方向トレース（安全目標 → ハザード）' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = ws1.max_row + 3
                ws1.cell(row=start_row - 1, column=1, value='逆方向トレース（SG → HZ）').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws1.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font

    # 列幅調整
    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws1.column_dimensions[col_letter].width = 18

    # --- シート2: SG→FSR ---
    ws2 = wb.create_sheet("SG→FSR")
    ws2['A1'] = 'SG → FSR トレーサビリティ（PH-04 確立）'
    ws2['A1'].font = title_font
    ws2.merge_cells('A1:F1')

    for section in sections:
        if '順方向トレース（安全目標 → 機能安全要求）' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = 3
                ws2.cell(row=start_row - 1, column=1, value='順方向トレース（SG → FSR）').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws2.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font
                            if '確立' in val:
                                cell.fill = pass_fill

        if '逆方向トレース（FSR → SG）' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = ws2.max_row + 3
                ws2.cell(row=start_row - 1, column=1, value='逆方向トレース（FSR → SG）').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws2.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font
                            if '展開済' in val:
                                cell.fill = pass_fill

    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws2.column_dimensions[col_letter].width = 20

    # --- シート3: FSR→TSR ---
    ws3 = wb.create_sheet("FSR→TSR")
    ws3['A1'] = 'FSR → TSR トレーサビリティ（PH-05 確立）'
    ws3['A1'].font = title_font
    ws3.merge_cells('A1:G1')

    for section in sections:
        if '順方向トレース（FSR → TSR）' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = 3
                ws3.cell(row=start_row - 1, column=1, value='順方向トレース（FSR → TSR）').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws3.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font
                            if '確立' in val:
                                cell.fill = pass_fill

        if '逆方向トレース（TSR → FSR → SG）' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = ws3.max_row + 3
                ws3.cell(row=start_row - 1, column=1, value='逆方向トレース（TSR → FSR → SG）').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws3.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font
                            if '展開待ち' in val:
                                cell.fill = tbd_fill

    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws3.column_dimensions[col_letter].width = 18

    # --- シート4: セキュリティトレース ---
    ws4 = wb.create_sheet("セキュリティ")
    ws4['A1'] = 'セキュリティトレース（段階的導入）'
    ws4['A1'].font = title_font
    ws4.merge_cells('A1:F1')

    for section in sections:
        if 'TH → SG 相互影響' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = 3
                ws4.cell(row=start_row - 1, column=1, value='TH → SG 相互影響マトリクス').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws4.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font
                            if 'TBD' in val:
                                cell.fill = tbd_fill

    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws4.column_dimensions[col_letter].width = 22

    # --- シート5: カバレッジサマリ ---
    ws5 = wb.create_sheet("カバレッジ")
    ws5['A1'] = 'トレーサビリティカバレッジサマリ'
    ws5['A1'].font = title_font
    ws5.merge_cells('A1:D1')

    for section in sections:
        if 'PH-05 時点のカバレッジ' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = 3
                ws5.cell(row=start_row - 1, column=1, value='PH-05 時点のカバレッジ').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws5.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font
                            if '**PASS**' in val or 'PASS' in val:
                                cell.fill = pass_fill
                            elif 'TBD' in val:
                                cell.fill = tbd_fill

        if 'フェーズ別展開計画' in section:
            table_lines = [l for l in section.split('\n') if l.strip().startswith('|')]
            rows = parse_markdown_table(table_lines)
            if rows:
                start_row = ws5.max_row + 3
                ws5.cell(row=start_row - 1, column=1, value='フェーズ別展開計画').font = subtitle_font
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        cell = ws5.cell(row=start_row + ri, column=ci + 1, value=val)
                        cell.border = thin_border
                        cell.alignment = wrap_align
                        if ri == 0:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                        else:
                            cell.font = normal_font
                            if '完了' in val:
                                cell.fill = pass_fill
                            elif '未着手' in val:
                                cell.fill = tbd_fill

    for col_letter in ['A', 'B', 'C', 'D']:
        ws5.column_dimensions[col_letter].width = 25

    wb.save(str(xlsx_path))
    print(f"  [OK] {xlsx_path.name}")


# ─── INDEX 生成 ──────────────────────────────────────

def generate_index_files():
    """各SPICEフォルダに 00_INDEX.md を生成"""
    count = 0
    for folder_name, info in sorted(FOLDER_INDEX.items()):
        folder_path = OUTPUT_BASE / folder_name
        folder_path.mkdir(parents=True, exist_ok=True)
        index_path = folder_path / "00_INDEX.md"

        # フォルダ内の既存ファイル一覧（INDEX自身を除く）
        existing = sorted(f.name for f in folder_path.iterdir()
                         if f.is_file() and f.name != "00_INDEX.md")

        lines = [
            f"# {folder_name}",
            f"",
            f"**プロセス**: {info['desc']}",
            f"**対応フェーズ**: {info['ph']}",
            f"",
            f"## 格納文書（計画）",
        ]
        for doc in info["docs"]:
            lines.append(f"- {doc}")
        lines.append(f"- プロセス記録 ([PR]_*.docx)")

        if existing:
            lines.append(f"")
            lines.append(f"## 現在のファイル")
            for f in existing:
                lines.append(f"- {f}")

        lines.append("")
        index_path.write_text("\n".join(lines), encoding="utf-8")
        count += 1

    print(f"[INDEX] 00_INDEX.md を {count} フォルダに生成")
    return count


# ─── ソースコードコピー ───────────────────────────────

SRC_DIR = BASE_DIR / "src"

def copy_source_code():
    """ソースコード (.c/.h) を SWE.3 フォルダにコピー"""
    dst_base = OUTPUT_BASE / "SWE.3_ソフトウェア詳細設計" / "src"
    count = 0
    if not SRC_DIR.exists():
        print("  [SKIP] src/ ディレクトリが見つかりません")
        return 0

    for src_file in sorted(SRC_DIR.rglob("*.[ch]")):
        rel = src_file.relative_to(SRC_DIR)
        dst_file = dst_base / rel
        dst_file.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src_file, dst_file)
        count += 1

    print(f"  [OK] ソースコード {count} ファイルを SWE.3/src/ にコピー")
    return count


def copy_project_json():
    """project.json を共通フォルダにコピー"""
    src = BASE_DIR / "project.json"
    dst = OUTPUT_BASE / "共通_プロジェクト管理" / "project.json"
    if src.exists():
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)
        print(f"  [OK] project.json をコピー")
        return 1
    return 0


# ─── メイン処理 ──────────────────────────────────────

def main():
    print(f"{'='*60}")
    print(f" 成果物エクスポート: {PROJECT_NAME}")
    print(f" 出力先: {OUTPUT_BASE}")
    print(f"{'='*60}")
    print()

    # 全SPICEフォルダ作成（設計成果物 + プロセス記録の全プロセス）
    all_folders = set(list(SPICE_FOLDERS.keys()) + list(PROCESS_RECORD_MAP.keys()))
    for folder_name in sorted(all_folders):
        folder_path = OUTPUT_BASE / folder_name
        folder_path.mkdir(parents=True, exist_ok=True)
        print(f"[FOLDER] {folder_name}/")

    print()

    # ファイル変換
    total = 0
    success = 0
    errors = []

    # ── 設計成果物（docs/） ──
    for folder_name, files in SPICE_FOLDERS.items():
        print(f"--- {folder_name} (設計成果物) ---")
        for src_name, dst_name in files:
            total += 1
            src_path = DOCS_DIR / src_name
            dst_path = OUTPUT_BASE / folder_name / dst_name

            if not src_path.exists():
                errors.append(f"  [SKIP] {src_name} — ファイルが見つかりません")
                print(f"  [SKIP] {src_name}")
                continue

            try:
                if dst_name.endswith('.xlsx'):
                    convert_trace_to_xlsx(src_path, dst_path)
                else:
                    convert_md_to_docx(src_path, dst_path)
                success += 1
            except Exception as e:
                errors.append(f"  [ERR] {src_name} → {dst_name}: {e}")
                print(f"  [ERR] {src_name}: {e}")

        print()

    # ── プロセス記録（process_records/）→ 各SPICEフォルダへ配置 ──
    print(f"--- プロセス記録 → SPICEフォルダ配置 ---")
    for folder_name, pr_filename in sorted(PROCESS_RECORD_MAP.items()):
        total += 1
        src_path = PR_DIR / pr_filename
        dst_name = "[PR]_" + pr_filename.replace('.md', '.docx')
        dst_path = OUTPUT_BASE / folder_name / dst_name

        if not src_path.exists():
            errors.append(f"  [SKIP] {pr_filename} — ファイルが見つかりません")
            print(f"  [SKIP] {pr_filename}")
            continue

        try:
            convert_md_to_docx(src_path, dst_path)
            success += 1
        except Exception as e:
            errors.append(f"  [ERR] {pr_filename}: {e}")
            print(f"  [ERR] {pr_filename}: {e}")
    print()

    # ── ソースコード ──
    print(f"--- SWE.3 ソースコード ---")
    src_count = copy_source_code()
    if src_count > 0:
        total += 1
        success += 1
    print()

    # ── project.json ──
    print(f"--- プロジェクト管理情報 ---")
    pj_count = copy_project_json()
    if pj_count > 0:
        total += 1
        success += 1
    print()

    # INDEX生成
    generate_index_files()
    print()

    # サマリ
    print(f"{'='*60}")
    print(f" 完了: {success}/{total} ファイル変換成功")
    if errors:
        print(f" エラー: {len(errors)} 件")
        for err in errors:
            print(f"   {err}")
    print(f" 出力先: {OUTPUT_BASE}")
    print(f"{'='*60}")

    return 0 if not errors else 1


if __name__ == '__main__':
    sys.exit(main())
