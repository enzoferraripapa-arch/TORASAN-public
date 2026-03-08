"""
TORASAN 仕様書ダイアグラム生成モジュール (DrawingML 版)

Word ネイティブ図形 (DrawingML XML) で5種のダイアグラムを生成する。
生成された図形はグループ化されており、Word 上でダブルクリックして個別編集可能。

依存: python-docx (parse_xml のみ使用)

方式: 全 XML を文字列で構築し parse_xml() で一括パース。
      lxml の etree.Element/tostring は使用しない (namespace prefix 問題回避)。
"""

from xml.sax.saxutils import escape as _esc
from docx.oxml.parser import parse_xml

from spec_constants import (
    SKILL_COUNT_UNIVERSAL,
    SKILL_COUNT_DOMAIN,
    SKILL_COUNT_TOTAL,
    KNOWLEDGE_COUNT_UNIVERSAL,
    KNOWLEDGE_COUNT_DOMAIN,
    KNOWLEDGE_COUNT_TOTAL,
    PHASE_COUNT,
)

# ============================================================
# カラーパレット (RGB tuples — spec_diagrams.py と統一)
# ============================================================
TERRACOTTA = (0xE0, 0x7A, 0x5F)
SAGE       = (0x4A, 0x8C, 0x6F)
SAND       = (0xF4, 0xF1, 0xDE)
CHARCOAL   = (0x2E, 0x30, 0x48)
MUSTARD    = (0xF2, 0xCC, 0x8F)
OLIVE      = (0x3D, 0x5C, 0x36)

LIGHT_SAGE  = (0xC8, 0xE0, 0xD2)
LIGHT_TERRA = (0xF0, 0xAA, 0x8F)
LIGHT_MUST  = (0xF9, 0xE4, 0xC0)
LIGHT_OLIVE = (0xA8, 0xBF, 0xA0)
LIGHT_CHAR  = (0x8A, 0x8D, 0xA0)

WHITE  = (0xFF, 0xFF, 0xFF)
BG     = (0xFA, 0xF9, 0xF6)
GRAY   = (0xBB, 0xBB, 0xBB)

# ============================================================
# XML 名前空間
# ============================================================
NS_WP  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
NS_A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
NS_WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
NS_WPG = 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup'
NS_W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

NS_DECLS = (
    f'xmlns:wp="{NS_WP}" xmlns:a="{NS_A}" '
    f'xmlns:wpg="{NS_WPG}" xmlns:wps="{NS_WPS}" xmlns:w="{NS_W}"'
)

# ============================================================
# ID 管理
# ============================================================
_shape_id_counter = 0

def _next_id():
    global _shape_id_counter
    _shape_id_counter += 1
    return _shape_id_counter

def reset_ids():
    global _shape_id_counter
    _shape_id_counter = 0

# ============================================================
# ユーティリティ
# ============================================================
EMU_PER_INCH = 914400

def rgb_hex(c):
    return '{:02X}{:02X}{:02X}'.format(*c)


_JC_MAP = {'ctr': 'center', 'l': 'left', 'r': 'right',
           'left': 'left', 'right': 'right', 'center': 'center', 'both': 'both'}


def _text_run_xml(text, font_size=16, bold=False, color=CHARCOAL,
                  font_name='Yu Gothic', align='ctr'):
    """w:p の XML 文字列を返す。
    font_size は half-points 単位 (例: 16 = 8pt, 20 = 10pt, 22 = 11pt)。"""
    jc_val = _JC_MAP.get(align, 'center')
    bold_xml = '<w:b/><w:bCs/>' if bold else ''
    return (
        f'<w:p>'
        f'<w:pPr><w:jc w:val="{jc_val}"/>'
        f'<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
        f'</w:pPr>'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:eastAsia="{font_name}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{font_size}"/><w:szCs w:val="{font_size}"/>'
        f'<w:color w:val="{rgb_hex(color)}"/>'
        f'{bold_xml}'
        f'</w:rPr>'
        f'<w:t>{_esc(text)}</w:t></w:r></w:p>'
    )


def _text_body_xml(texts, font_name='Yu Gothic', align='ctr',
                   font_size=16, bold=False, color=CHARCOAL,
                   v_anchor='ctr'):
    """wps:txbx + wps:bodyPr の XML 文字列を返す"""
    if isinstance(texts, str):
        texts = [texts]

    paras = []
    for item in texts:
        if isinstance(item, str):
            txt, ov = item, {}
        else:
            txt, ov = item
        paras.append(_text_run_xml(
            txt,
            font_size=ov.get('font_size', font_size),
            bold=ov.get('bold', bold),
            color=ov.get('color', color),
            font_name=ov.get('font_name', font_name),
            align=ov.get('align', align),
        ))

    content = ''.join(paras)
    return (
        f'<wps:txbx><w:txbxContent>{content}</w:txbxContent></wps:txbx>'
        f'<wps:bodyPr anchor="{v_anchor}" anchorCtr="0"'
        f' lIns="36000" rIns="36000" tIns="18000" bIns="18000">'
        f'<a:normAutofit/>'
        f'</wps:bodyPr>'
    )


# ============================================================
# プリミティブビルダー (全て XML 文字列を返す)
# ============================================================

def make_shape(x, y, cx, cy, preset, fill=None, outline=None,
               line_w=12700, texts=None, text_props=None, radius_adj=None,
               no_fill=False, no_line=False):
    """XML 文字列を返す wps:wsp"""
    # geometry
    adj_xml = ''
    if radius_adj is not None:
        adj_xml = f'<a:gd name="adj" fmla="val {radius_adj}"/>'

    # fill
    if no_fill:
        fill_xml = '<a:noFill/>'
    elif fill:
        fill_xml = f'<a:solidFill><a:srgbClr val="{rgb_hex(fill)}"/></a:solidFill>'
    else:
        fill_xml = ''

    # line
    if no_line:
        ln_xml = f'<a:ln w="{int(line_w)}"><a:noFill/></a:ln>'
    elif outline:
        ln_xml = (f'<a:ln w="{int(line_w)}">'
                  f'<a:solidFill><a:srgbClr val="{rgb_hex(outline)}"/></a:solidFill>'
                  f'</a:ln>')
    else:
        ln_xml = f'<a:ln w="{int(line_w)}"/>'

    # text
    if texts is not None:
        tp = text_props or {}
        body_xml = _text_body_xml(texts, **tp)
    else:
        body_xml = '<wps:bodyPr/>'

    return (
        f'<wps:wsp>'
        f'<wps:cNvSpPr/>'
        f'<wps:spPr>'
        f'<a:xfrm><a:off x="{int(x)}" y="{int(y)}"/>'
        f'<a:ext cx="{int(cx)}" cy="{int(cy)}"/></a:xfrm>'
        f'<a:prstGeom prst="{preset}"><a:avLst>{adj_xml}</a:avLst></a:prstGeom>'
        f'{fill_xml}{ln_xml}'
        f'</wps:spPr>'
        f'{body_xml}'
        f'</wps:wsp>'
    )


def make_connector(x1, y1, x2, y2, color=CHARCOAL, line_w=12700,
                   arrow_head=False, dashed=False):
    """XML 文字列を返す wps:wsp (コネクタ)"""
    off_x = min(x1, x2)
    off_y = min(y1, y2)
    dx = abs(x2 - x1)
    dy = abs(y2 - y1)

    flip_attrs = ''
    if x2 < x1:
        flip_attrs += ' flipH="1"'
    if y2 < y1:
        flip_attrs += ' flipV="1"'

    dash_xml = '<a:prstDash val="dash"/>' if dashed else ''
    arrow_xml = '<a:tailEnd type="triangle" w="med" len="med"/>' if arrow_head else ''

    return (
        f'<wps:wsp>'
        f'<wps:cNvCnPr/>'
        f'<wps:spPr>'
        f'<a:xfrm{flip_attrs}>'
        f'<a:off x="{int(off_x)}" y="{int(off_y)}"/>'
        f'<a:ext cx="{max(int(dx), 1)}" cy="{max(int(dy), 1)}"/>'
        f'</a:xfrm>'
        f'<a:prstGeom prst="straightConnector1"><a:avLst/></a:prstGeom>'
        f'<a:ln w="{int(line_w)}">'
        f'<a:solidFill><a:srgbClr val="{rgb_hex(color)}"/></a:solidFill>'
        f'{dash_xml}{arrow_xml}'
        f'</a:ln>'
        f'</wps:spPr>'
        f'<wps:bodyPr/>'
        f'</wps:wsp>'
    )


def make_text_label(x, y, cx, cy, text, font_size=20, bold=False,
                    color=CHARCOAL, align='ctr', v_anchor='ctr'):
    return make_shape(x, y, cx, cy, 'rect',
                      no_fill=True, no_line=True,
                      texts=text,
                      text_props={
                          'font_size': font_size, 'bold': bold,
                          'color': color, 'align': align, 'v_anchor': v_anchor,
                      })


def make_group(logical_w, logical_h, physical_w_emu, physical_h_emu, children):
    """children は XML 文字列のリスト。wpg:wgp 全体の XML 文字列を返す。"""
    inner = ''.join(children)
    return (
        f'<wpg:wgp>'
        f'<wpg:cNvGrpSpPr/>'
        f'<wpg:grpSpPr>'
        f'<a:xfrm>'
        f'<a:off x="0" y="0"/>'
        f'<a:ext cx="{physical_w_emu}" cy="{physical_h_emu}"/>'
        f'<a:chOff x="0" y="0"/>'
        f'<a:chExt cx="{int(logical_w)}" cy="{int(logical_h)}"/>'
        f'</a:xfrm>'
        f'</wpg:grpSpPr>'
        f'{inner}'
        f'</wpg:wgp>'
    )


def make_inline(physical_w_emu, physical_h_emu, name, wgp_xml):
    """wp:inline を parse_xml() で生成。呼び出し側で run._r.add_drawing(inline)。"""
    shape_id = _next_id()
    xml = (
        f'<wp:inline {NS_DECLS}'
        f' distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{physical_w_emu}" cy="{physical_h_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{shape_id}" name="{_esc(name)}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic><a:graphicData uri="{NS_WPG}">'
        f'{wgp_xml}'
        f'</a:graphicData></a:graphic>'
        f'</wp:inline>'
    )
    return parse_xml(xml)


make_drawing = make_inline


def _diagram_emu(w_inches, h_inches):
    """図の物理サイズ (インチ) を EMU に変換して返す。
    _Y_SCALE に頼らず、図ごとに適切なサイズを明示指定する。"""
    return int(w_inches * EMU_PER_INCH), int(h_inches * EMU_PER_INCH)


# ============================================================
# 図1: リポジトリ構成図 (S2.1)
# ============================================================
def draw_repo_structure():
    LW, LH = 2400, 1200
    PW, PH = _diagram_emu(6.0, 4.5)
    shapes = []

    shapes.append(make_shape(0, 0, LW, LH, 'rect', fill=BG, no_line=True))

    shapes.append(make_text_label(LW//2 - 400, 10, 800, 60,
                                  'TORASAN リポジトリ構成',
                                  font_size=17, bold=True, color=CHARCOAL))

    root_x, root_y = LW//2, 110
    root_w, root_h = 340, 65
    shapes.append(make_shape(root_x - root_w//2, root_y - root_h//2, root_w, root_h,
                             'roundRect', fill=CHARCOAL, outline=CHARCOAL,
                             radius_adj=10000,
                             texts='TORASAN/',
                             text_props={'font_size': 22, 'bold': True, 'color': WHITE}))

    children = [
        ('.claude/skills/', f'{SKILL_COUNT_TOTAL} skills', SAGE, LIGHT_SAGE),
        ('.claude/knowledge/', f'{KNOWLEDGE_COUNT_TOTAL} files', TERRACOTTA, LIGHT_TERRA),
        ('app/', 'React + Fastify', MUSTARD, LIGHT_MUST),
        ('scripts/', 'ツール・生成', OLIVE, LIGHT_OLIVE),
        ('PROCESS.md', 'V-model 定義', CHARCOAL, LIGHT_CHAR),
        ('project.json', 'SoT', CHARCOAL, LIGHT_CHAR),
        ('packages/', '規格テンプレート', TERRACOTTA, LIGHT_TERRA),
        ('docs/', '成果物文書', SAGE, LIGHT_SAGE),
    ]

    start_y = 240
    row_h = 110
    box_w, box_h = 340, 85
    col_x = [LW//4 + 50, 3*LW//4 - 50]

    shapes.append(make_connector(root_x, root_y + root_h//2, root_x, start_y - 30,
                                 color=GRAY, line_w=12700))
    shapes.append(make_connector(col_x[0], start_y - 30, col_x[1], start_y - 30,
                                 color=GRAY, line_w=12700))

    for i, (label, desc, color, light) in enumerate(children):
        col = i % 2
        row = i // 2
        cx = col_x[col]
        cy = start_y + row * row_h

        shapes.append(make_connector(cx, start_y - 30, cx, cy - box_h//2,
                                     color=GRAY, line_w=12700))
        shapes.append(make_shape(
            cx - box_w//2, cy - box_h//2, box_w, box_h, 'roundRect',
            fill=light, outline=color, line_w=12700, radius_adj=8000,
            texts=[(label, {'font_size': 20, 'bold': True, 'color': color, 'align': 'left'}),
                   (desc, {'font_size': 16, 'color': CHARCOAL, 'align': 'left'})],
            text_props={'v_anchor': 'ctr'},
        ))

    legend_y = LH - 70
    legend_items = [
        ('Skills', SAGE), ('Knowledge', TERRACOTTA),
        ('App', MUSTARD), ('Tools', OLIVE), ('Config', CHARCOAL),
    ]
    lx = 200
    for lbl, c in legend_items:
        shapes.append(make_shape(lx, legend_y, 22, 22, 'roundRect',
                                 fill=c, outline=c, radius_adj=5000))
        shapes.append(make_text_label(lx + 28, legend_y - 3, 130, 28, lbl,
                                      font_size=16, color=CHARCOAL, align='left'))
        lx += 180

    wgp = make_group(LW, LH, PW, PH, shapes)
    return make_drawing(PW, PH, 'リポジトリ構成図', wgp)


# ============================================================
# 図2: 二層配布モデル (S2.2)
# ============================================================
def draw_distribution_model():
    LW, LH = 2400, 1400
    PW, PH = _diagram_emu(6.0, 5.0)
    shapes = []

    shapes.append(make_shape(0, 0, LW, LH, 'rect', fill=BG, no_line=True))

    shapes.append(make_text_label(LW//2 - 300, 10, 600, 50,
                                  '二層配布モデル',
                                  font_size=17, bold=True, color=CHARCOAL))

    hub_x, hub_y = LW//2, 200
    hub_w, hub_h = 500, 80
    shapes.append(make_shape(
        hub_x - hub_w//2, hub_y - hub_h//2, hub_w, hub_h, 'roundRect',
        fill=TERRACOTTA, outline=TERRACOTTA, radius_adj=10000,
        texts=[('TORASAN リポジトリ',
                {'font_size': 22, 'bold': True, 'color': WHITE}),
               (f'Skills: {SKILL_COUNT_TOTAL} / Knowledge: {KNOWLEDGE_COUNT_TOTAL}',
                {'font_size': 16, 'color': (255, 230, 220)})],
        text_props={'v_anchor': 'ctr'},
    ))

    bw, bh = 320, 75

    # Tier 1
    t1_x = LW // 4
    shapes.append(make_text_label(t1_x - 100, 295, 200, 35,
                                  'Tier 1: 汎用', font_size=20, bold=True, color=SAGE))
    t1_cmd_y = 380
    shapes.append(make_shape(t1_x - bw//2, t1_cmd_y - bh//2, bw, bh, 'roundRect',
                             fill=SAGE, outline=SAGE, radius_adj=8000,
                             texts='install.sh',
                             text_props={'font_size': 22, 'bold': True, 'color': WHITE}))
    shapes.append(make_connector(hub_x - 100, hub_y + hub_h//2,
                                 t1_x, t1_cmd_y - bh//2,
                                 color=SAGE, arrow_head=True, line_w=19050))

    t1_dest_y = 560
    shapes.append(make_shape(
        t1_x - bw//2, t1_dest_y - bh//2, bw, bh, 'roundRect',
        fill=LIGHT_SAGE, outline=SAGE, line_w=12700, radius_adj=8000,
        texts=[('~/.claude/', {'font_size': 22, 'bold': True, 'color': SAGE}),
               (f'skills/ ({SKILL_COUNT_UNIVERSAL}) + knowledge/ ({KNOWLEDGE_COUNT_UNIVERSAL})',
                {'font_size': 16, 'color': CHARCOAL})],
        text_props={'v_anchor': 'ctr'},
    ))
    shapes.append(make_connector(t1_x, t1_cmd_y + bh//2, t1_x, t1_dest_y - bh//2,
                                 color=SAGE, arrow_head=True, line_w=19050))

    t1_tgt_y = 740
    shapes.append(make_shape(t1_x - 140, t1_tgt_y - 25, 280, 50, 'roundRect',
                             fill=WHITE, outline=SAGE, line_w=12700, radius_adj=6000,
                             texts='全プロジェクト共通',
                             text_props={'font_size': 17, 'color': CHARCOAL}))
    shapes.append(make_connector(t1_x, t1_dest_y + bh//2, t1_x, t1_tgt_y - 25,
                                 color=SAGE, dashed=True, line_w=12700))

    # Tier 2
    t2_x = 3 * LW // 4
    shapes.append(make_text_label(t2_x - 120, 295, 240, 35,
                                  'Tier 2: ドメイン', font_size=20, bold=True, color=TERRACOTTA))
    t2_cmd_y = 380
    shapes.append(make_shape(t2_x - bw//2, t2_cmd_y - bh//2, bw, bh, 'roundRect',
                             fill=TERRACOTTA, outline=TERRACOTTA, radius_adj=8000,
                             texts='/repo-manage sync',
                             text_props={'font_size': 22, 'bold': True, 'color': WHITE}))
    shapes.append(make_connector(hub_x + 100, hub_y + hub_h//2,
                                 t2_x, t2_cmd_y - bh//2,
                                 color=TERRACOTTA, arrow_head=True, line_w=19050))

    t2_dest_y = 560
    shapes.append(make_shape(
        t2_x - bw//2, t2_dest_y - bh//2, bw, bh, 'roundRect',
        fill=LIGHT_TERRA, outline=TERRACOTTA, line_w=12700, radius_adj=8000,
        texts=[('PJ/.claude/skills/', {'font_size': 22, 'bold': True, 'color': TERRACOTTA}),
               (f'{SKILL_COUNT_DOMAIN} ドメインスキル（選択配布）',
                {'font_size': 16, 'color': CHARCOAL})],
        text_props={'v_anchor': 'ctr'},
    ))
    shapes.append(make_connector(t2_x, t2_cmd_y + bh//2, t2_x, t2_dest_y - bh//2,
                                 color=TERRACOTTA, arrow_head=True, line_w=19050))

    pjs = ['PJ-A (自動車)', 'PJ-B (家電)', 'PJ-C (産業)']
    t2_tgt_y = 720
    for i, pj in enumerate(pjs):
        py = t2_tgt_y + i * 65
        shapes.append(make_shape(t2_x - 140, py - 25, 280, 50, 'roundRect',
                                 fill=WHITE, outline=TERRACOTTA, line_w=12700, radius_adj=6000,
                                 texts=pj,
                                 text_props={'font_size': 17, 'color': CHARCOAL}))
        shapes.append(make_connector(t2_x, t2_dest_y + bh//2, t2_x, py - 25,
                                     color=TERRACOTTA, dashed=True, line_w=12700))

    shapes.append(make_connector(LW//2, 270, LW//2, LH - 80,
                                 color=GRAY, dashed=True, line_w=6350))

    wgp = make_group(LW, LH, PW, PH, shapes)
    return make_drawing(PW, PH, '二層配布モデル', wgp)


# ============================================================
# 図3: メモリ・状態管理アーキテクチャ (S2.3)
# ============================================================
def draw_memory_architecture():
    LW, LH = 2400, 1400
    PW, PH = _diagram_emu(6.0, 5.0)
    shapes = []

    shapes.append(make_shape(0, 0, LW, LH, 'rect', fill=BG, no_line=True))

    shapes.append(make_text_label(LW//2 - 500, 10, 1000, 55,
                                  'メモリ・状態管理アーキテクチャ',
                                  font_size=17, bold=True, color=CHARCOAL))

    zones = [
        ('Session スコープ（揮発性）', TERRACOTTA, LIGHT_TERRA,
         [('Auto Memory', 'MEMORY.md', '自動ロード (200行上限)'),
          ('Session State', 'session_state.md', '/session end で上書き')]),
        ('Project スコープ（Git 永続）', SAGE, LIGHT_SAGE,
         [('Project Knowledge', '.claude/knowledge/*.md', 'ドメイン知識'),
          ('Process Records', 'process_records/*.md', 'SPICE エビデンス'),
          ('project.json', 'SoT', '構成・進捗・変更ログ')]),
        ('Global スコープ（永続）', MUSTARD, LIGHT_MUST,
         [('Project Registry', 'project_registry.json', 'PJ 一覧'),
          ('Skill Manifest', '.shared-skills-manifest.json', '配布履歴')]),
    ]

    margin_x = 60
    zone_w = LW - 2 * margin_x
    y_cursor = 90

    for zone_label, zone_color, zone_bg, items in zones:
        zone_h = 70 + len(items) * 85

        shapes.append(make_shape(margin_x, y_cursor, zone_w, zone_h, 'roundRect',
                                 fill=zone_bg, outline=zone_color, line_w=19050,
                                 radius_adj=5000))

        shapes.append(make_text_label(margin_x + 20, y_cursor + 5, 600, 40,
                                      zone_label, font_size=18, bold=True,
                                      color=zone_color, align='left'))

        item_y = y_cursor + 60
        item_w = 650
        item_h = 80
        cols = min(len(items), 3)
        gap = 30
        total_w = cols * item_w + (cols - 1) * gap
        start_x = margin_x + (zone_w - total_w) // 2

        for j, (name, path, desc) in enumerate(items):
            ix = start_x + j * (item_w + gap)
            iy = item_y
            shapes.append(make_shape(
                ix, iy, item_w, item_h, 'roundRect',
                fill=WHITE, outline=zone_color, line_w=12700, radius_adj=6000,
                texts=[(name, {'font_size': 20, 'bold': True, 'color': zone_color, 'align': 'left'}),
                       (f'{path} — {desc}', {'font_size': 15, 'color': CHARCOAL, 'align': 'left'})],
                text_props={'v_anchor': 'ctr'},
            ))

        y_cursor += zone_h + 25

    arrow_x = LW - 45
    shapes.append(make_connector(arrow_x, 130, arrow_x, y_cursor - 40,
                                 color=CHARCOAL, arrow_head=True, line_w=12700))
    shapes.append(make_text_label(arrow_x - 25, 90, 50, 35,
                                  '揮', font_size=16, color=CHARCOAL))
    shapes.append(make_text_label(arrow_x - 25, y_cursor - 65, 50, 35,
                                  '永', font_size=16, color=CHARCOAL))

    wgp = make_group(LW, LH, PW, PH, shapes)
    return make_drawing(PW, PH, 'メモリ・状態管理アーキテクチャ', wgp)


# ============================================================
# 図4: V モデル (S5.1)
# ============================================================
def draw_vmodel():
    LW, LH = 3600, 2400
    PW, PH = _diagram_emu(6.5, 6.0)
    shapes = []

    shapes.append(make_shape(0, 0, LW, LH, 'rect', fill=BG, no_line=True))

    shapes.append(make_text_label(LW//2 - 600, 10, 1200, 60,
                                  f'V モデル — {PHASE_COUNT} フェーズ構成',
                                  font_size=26, bold=True, color=CHARCOAL))

    CAT_COLORS = {
        'MAN': (CHARCOAL, LIGHT_CHAR),
        'SYS': (SAGE, LIGHT_SAGE),
        'SWE': (TERRACOTTA, LIGHT_TERRA),
        'SUP': (OLIVE, LIGHT_OLIVE),
    }

    left_phases = [
        ('PH-01', '計画・安全計画', 'MAN'),
        ('PH-02', 'アイテム定義', 'SYS'),
        ('PH-03', 'HARA', 'SYS'),
        ('PH-04', 'FSC', 'SYS'),
        ('PH-05', 'TSC', 'SYS'),
        ('PH-06', 'システム設計', 'SYS'),
        ('PH-07', 'HW 設計', 'SYS'),
        ('PH-08', 'SRS（SW安全要求）', 'SWE'),
    ]
    right_phases = [
        ('PH-15', '安全アセスメント', 'SUP'),
        ('PH-14', '安全検証', 'SYS'),
        ('PH-13', 'システムテスト', 'SYS'),
        ('PH-12', 'HW 統合テスト', 'SYS'),
        ('PH-11', 'SW テスト', 'SWE'),
        ('PH-10', 'SW ユニット', 'SWE'),
        ('PH-09', 'SW アーキ設計', 'SWE'),
    ]

    n_rows = 8
    y_top = 120
    y_bottom = 2200
    y_step = (y_bottom - y_top) / (n_rows - 1)

    x_left_start = 350
    x_right_start = 3250
    x_center = LW // 2
    x_indent = (x_center - x_left_start) / (n_rows - 1)

    box_w, box_h = 400, 100

    def _phase_box(cx, cy, pid, name, cat):
        c_dark, c_light = CAT_COLORS[cat]
        x0 = cx - box_w // 2
        y0 = cy - box_h // 2
        shapes.append(make_shape(
            x0, y0, box_w, box_h, 'roundRect',
            fill=c_light, outline=c_dark, line_w=12700, radius_adj=8000,
            texts=[(f'{pid}  {name}',
                    {'font_size': 18, 'bold': True, 'color': c_dark, 'align': 'ctr'})],
            text_props={'v_anchor': 'ctr'},
        ))
        return (cx, cy)

    left_centers = []
    for i, (pid, name, cat) in enumerate(left_phases):
        cx = x_center if i == 7 else x_left_start + i * x_indent
        cy = y_top + i * y_step
        left_centers.append(_phase_box(cx, cy, pid, name, cat))

    right_centers = []
    for i, (pid, name, cat) in enumerate(right_phases):
        cx = x_right_start - i * x_indent
        cy = y_top + i * y_step
        right_centers.append(_phase_box(cx, cy, pid, name, cat))

    for i in range(len(left_centers) - 1):
        ax, ay = left_centers[i]
        bx, by = left_centers[i + 1]
        shapes.append(make_connector(ax, ay + box_h//2, bx, by - box_h//2,
                                     color=CHARCOAL, line_w=12700))

    for i in range(len(right_centers) - 1):
        ax, ay = right_centers[i]
        bx, by = right_centers[i + 1]
        shapes.append(make_connector(ax, ay + box_h//2, bx, by - box_h//2,
                                     color=CHARCOAL, line_w=12700))

    ph08 = left_centers[7]
    ph09 = right_centers[6]
    shapes.append(make_connector(ph08[0] + box_w//2, ph08[1],
                                 ph09[0] - box_w//2, ph09[1],
                                 color=CHARCOAL, arrow_head=True, line_w=12700))

    for i in range(7):
        lx, ly = left_centers[i]
        rx, ry = right_centers[i]
        shapes.append(make_connector(lx + box_w//2 + 10, ly,
                                     rx - box_w//2 - 10, ry,
                                     color=GRAY, dashed=True, line_w=6350))

    ly = LH - 90
    lx = 200
    for label, cat_key in [('MAN (管理)', 'MAN'), ('SYS (システム)', 'SYS'),
                           ('SWE (ソフトウェア)', 'SWE'), ('SUP (サポート)', 'SUP')]:
        c_dark, c_light = CAT_COLORS[cat_key]
        shapes.append(make_shape(lx, ly, 30, 24, 'roundRect',
                                 fill=c_light, outline=c_dark, line_w=12700, radius_adj=5000))
        shapes.append(make_text_label(lx + 38, ly - 3, 280, 30, label,
                                      font_size=18, color=CHARCOAL, align='left'))
        lx += 360

    dlx = lx + 40
    shapes.append(make_connector(dlx, ly + 12, dlx + 40, ly + 12,
                                 color=GRAY, dashed=True, line_w=6350))
    shapes.append(make_text_label(dlx + 48, ly - 3, 200, 30,
                                  '検証対応関係', font_size=18, color=GRAY, align='left'))

    wgp = make_group(LW, LH, PW, PH, shapes)
    return make_drawing(PW, PH, 'V モデル 15 フェーズ構成', wgp)


# ============================================================
# 図5: ツールパス解決フロー (S8.4)
# ============================================================
def draw_tool_resolution():
    LW, LH = 2000, 1400
    PW, PH = _diagram_emu(5.5, 5.0)
    shapes = []

    shapes.append(make_shape(0, 0, LW, LH, 'rect', fill=BG, no_line=True))

    shapes.append(make_text_label(LW//2 - 350, 5, 700, 50,
                                  'ツールパス解決フロー',
                                  font_size=20, bold=True, color=CHARCOAL))

    cx = LW // 2

    sy = 100
    shapes.append(make_shape(cx - 120, sy - 25, 240, 50, 'flowChartTerminator',
                             fill=CHARCOAL, outline=CHARCOAL,
                             texts='ツール実行要求',
                             text_props={'font_size': 20, 'bold': True, 'color': WHITE}))

    d_w, d_h = 320, 120

    def _decision_block(dy, q1, q2, yes_label, yes_x):
        shapes.append(make_shape(cx - d_w//2, dy - d_h//2, d_w, d_h, 'diamond',
                                 fill=LIGHT_MUST, outline=MUSTARD, line_w=12700,
                                 texts=[(q1, {'font_size': 20, 'bold': True, 'color': CHARCOAL}),
                                        (q2, {'font_size': 15, 'color': CHARCOAL})],
                                 text_props={'v_anchor': 'ctr'}))
        shapes.append(make_shape(yes_x - 130, dy - 30, 260, 60, 'roundRect',
                                 fill=LIGHT_SAGE, outline=SAGE, line_w=12700, radius_adj=8000,
                                 texts=yes_label,
                                 text_props={'font_size': 20, 'bold': True, 'color': SAGE}))
        shapes.append(make_connector(cx + d_w//2, dy, yes_x - 130, dy,
                                     color=SAGE, arrow_head=True, line_w=12700))
        shapes.append(make_text_label(cx + d_w//2 + 5, dy - 30, 50, 25,
                                      'Yes', font_size=18, bold=True, color=SAGE, align='left'))
        shapes.append(make_text_label(cx + 8, dy + d_h//2 + 2, 40, 25,
                                      'No', font_size=18, bold=True, color=TERRACOTTA, align='left'))

    yes_x = cx + 380

    d1y = 250
    shapes.append(make_connector(cx, sy + 25, cx, d1y - d_h//2,
                                 color=CHARCOAL, arrow_head=True, line_w=12700))
    _decision_block(d1y, 'cmd.exe PATH', 'に存在？', '直接実行', yes_x)

    d2y = 480
    shapes.append(make_connector(cx, d1y + d_h//2, cx, d2y - d_h//2,
                                 color=CHARCOAL, arrow_head=True, line_w=12700))
    _decision_block(d2y, 'KNOWN_TOOL_PATHS', 'に登録？', 'フルパス実行', yes_x)

    d3y = 710
    shapes.append(make_connector(cx, d2y + d_h//2, cx, d3y - d_h//2,
                                 color=CHARCOAL, arrow_head=True, line_w=12700))
    _decision_block(d3y, 'Git Bash', 'で発見？', 'Git Bash 経由', yes_x)

    fail_y = 910
    shapes.append(make_connector(cx, d3y + d_h//2, cx, fail_y - 30,
                                 color=TERRACOTTA, arrow_head=True, line_w=12700))
    shapes.append(make_shape(cx - 150, fail_y - 30, 300, 60, 'roundRect',
                             fill=LIGHT_TERRA, outline=TERRACOTTA, line_w=12700, radius_adj=8000,
                             texts='ツール未検出 (missing)',
                             text_props={'font_size': 20, 'bold': True, 'color': TERRACOTTA}))

    end_y = 1050
    end_x = yes_x
    for yy in [d1y, d2y, d3y]:
        shapes.append(make_connector(end_x, yy + 30, end_x, end_y - 35,
                                     color=SAGE, line_w=12700))

    shapes.append(make_shape(end_x - 120, end_y - 30, 240, 60, 'flowChartTerminator',
                             fill=SAGE, outline=SAGE,
                             texts='実行完了',
                             text_props={'font_size': 20, 'bold': True, 'color': WHITE}))

    note_y = 1150
    note_x = LW // 2
    shapes.append(make_shape(
        note_x - 400, note_y, 800, 120, 'roundRect',
        fill=WHITE, outline=GRAY, line_w=6350, radius_adj=5000,
        texts=[('解決順序: cmd.exe PATH → KNOWN_TOOL_PATHS → Git Bash',
                {'font_size': 18, 'bold': True, 'color': CHARCOAL, 'align': 'left'}),
               ('KNOWN_TOOL_PATHS: environmentService.ts で管理',
                {'font_size': 15, 'color': CHARCOAL, 'align': 'left'}),
               ('新規ツール追加時: error_prevention.md §E フロー参照',
                {'font_size': 15, 'color': CHARCOAL, 'align': 'left'})],
        text_props={'v_anchor': 't'},
    ))

    wgp = make_group(LW, LH, PW, PH, shapes)
    return make_drawing(PW, PH, 'ツールパス解決フロー', wgp)


# ============================================================
# テスト用エントリポイント
# ============================================================
if __name__ == '__main__':
    import os
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    reset_ids()
    doc = Document()

    funcs = [
        ('リポジトリ構成図', draw_repo_structure),
        ('二層配布モデル', draw_distribution_model),
        ('メモリ・状態管理アーキテクチャ', draw_memory_architecture),
        ('V モデル', draw_vmodel),
        ('ツールパス解決フロー', draw_tool_resolution),
    ]

    for title, func in funcs:
        doc.add_heading(title, level=2)
        inline = func()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run._r.add_drawing(inline)
        doc.add_page_break()

    out_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'diagrams')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, 'test_drawingml.docx')
    doc.save(out_path)
    print(f'Test DOCX saved: {out_path}')
