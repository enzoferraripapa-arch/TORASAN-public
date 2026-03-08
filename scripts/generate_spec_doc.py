"""TORASAN フレームワーク仕様書 Word文書生成スクリプト"""

import os
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# 共有定数・ダイアグラムモジュール
sys.path.insert(0, os.path.dirname(__file__))
from spec_constants import (
    SKILL_COUNT_UNIVERSAL,
    SKILL_COUNT_DOMAIN,
    SKILL_COUNT_TOTAL,
    KNOWLEDGE_COUNT_UNIVERSAL,
    KNOWLEDGE_COUNT_DOMAIN,
    KNOWLEDGE_COUNT_TOTAL,
    PHASE_COUNT,
    DIAGRAM_MODE,
    SPEC_INITIAL_DATE,
    FRAMEWORK_VERSION,
    PROCESS_VERSION,
)
from spec_diagrams import (
    draw_repo_structure,
    draw_distribution_model,
    draw_memory_architecture,
    draw_vmodel,
    draw_tool_resolution,
)

# DrawingML 版（編集可能図形）— モード切替
_dml_funcs = {}
if DIAGRAM_MODE == 'drawingml':
    try:
        from spec_diagrams_drawingml import (
            draw_repo_structure as _dml_repo,
            draw_distribution_model as _dml_dist,
            draw_memory_architecture as _dml_mem,
            draw_vmodel as _dml_vmodel,
            draw_tool_resolution as _dml_tool,
            reset_ids,
        )
        _dml_funcs = {
            'repo': _dml_repo,
            'dist': _dml_dist,
            'mem': _dml_mem,
            'vmodel': _dml_vmodel,
            'tool': _dml_tool,
        }
    except ImportError:
        pass  # DrawingML モジュールが無い場合は PNG フォールバック


def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shading_elm)


def _insert_diagram(doc, diagram_func, width=Inches(6), caption=None,
                    dml_key=None):
    """ダイアグラムを DOCX に挿入（DrawingML 優先、PNG フォールバック）

    dml_key: _dml_funcs のキー（'repo', 'dist' 等）。指定時は DrawingML 版を優先使用。
    """
    def _add_caption(cap_text):
        cap = doc.add_paragraph(cap_text)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    try:
        dml_func = _dml_funcs.get(dml_key) if dml_key else None
        if dml_func:
            inline = dml_func()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run._r.add_drawing(inline)
        else:
            buf = diagram_func()
            doc.add_picture(buf, width=width)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            _add_caption(caption)
    except Exception as e:
        # DrawingML 失敗時は PNG フォールバック
        if dml_key and _dml_funcs.get(dml_key):
            try:
                buf = diagram_func()
                doc.add_picture(buf, width=width)
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    _add_caption(f'{caption} [PNG fallback]')
                return
            except Exception:
                pass
        p = doc.add_paragraph(f'[図の生成に失敗: {e}]')
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


def add_styled_table(doc, headers, rows, col_widths=None):
    """スタイル付きテーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # ヘッダー行
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2E4057')

    # データ行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def create_spec_doc(output_path, version_label='v1.0', revision_history=None):
    # DrawingML ID カウンタリセット
    if _dml_funcs:
        reset_ids()

    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # フォント設定
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Yu Gothic'
        h_style.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # ========== 表紙 ==========
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TORASAN')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('機能安全開発プロセスフレームワーク')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    doc.add_paragraph()

    doc_title = doc.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_title.add_run('フレームワーク仕様書')
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for _ in range(4):
        doc.add_paragraph()

    # 文書情報テーブル
    info_table = doc.add_table(rows=6, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'
    info_data = [
        ('文書番号', 'TORASAN-SPEC-001'),
        ('版数', version_label),
        ('初版作成日', SPEC_INITIAL_DATE),
        ('最終更新日', date.today().strftime('%Y-%m-%d')),
        ('フレームワーク版数', FRAMEWORK_VERSION),
        ('プロセス版数', PROCESS_VERSION),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        set_cell_shading(info_table.rows[i].cells[0], 'E8EDF2')
        for cell in info_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_page_break()

    # ========== 改版履歴 ==========
    doc.add_heading('改版履歴', level=1)
    rev_rows = revision_history if revision_history else [
        ['v1.0', SPEC_INITIAL_DATE, '初版作成', 'TORA'],
    ]
    add_styled_table(doc,
        ['版数', '日付', '変更内容', '承認'],
        rev_rows,
        col_widths=[2, 3, 8, 3]
    )
    doc.add_page_break()

    # ========== 目次プレースホルダ ==========
    doc.add_heading('目次', level=1)
    p = doc.add_paragraph('（Word の「参考資料」→「目次の更新」で自動生成してください）')
    p.italic = True
    doc.add_page_break()

    # ========== 1. 概要 ==========
    doc.add_heading('1. 概要', level=1)

    doc.add_heading('1.1 目的', level=2)
    doc.add_paragraph(
        'TORASAN（Toolkit for Organized Risk Analysis and Safety Assurance Navigation）は、'
        'ISO 26262 および Automotive SPICE に基づく機能安全開発プロセスを、'
        'Claude Code のスキルシステムで自動化・標準化するフレームワークである。'
    )
    doc.add_paragraph(
        '本文書は TORASAN フレームワークの技術仕様を定義し、'
        'フレームワークの構成要素・動作仕様・データ構造・プロセス定義を網羅的に記述する。'
    )

    doc.add_heading('1.2 適用範囲', level=2)
    doc.add_paragraph('本仕様書は以下を対象とする:')
    items = [
        'TORASAN フレームワーク本体（スキル・ナレッジ・プロセス定義）',
        'フレームワークを利用する全プロジェクト（機能安全・一般開発）',
        'スキル配布・同期メカニズム',
        'セッション管理・状態保存メカニズム',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('1.3 用語定義', level=2)
    add_styled_table(doc,
        ['用語', '定義'],
        [
            ['TORASAN', '機能安全開発プロセスフレームワークの名称'],
            ['スキル', 'Claude Code のカスタムスラッシュコマンド（/skill-name 形式）'],
            ['ナレッジ', 'スキル実行時に参照されるドメイン知識ファイル（.md 形式）'],
            ['ワークツリー', 'Git worktree による並列開発ブランチ環境'],
            ['SPICE', 'Automotive SPICE — 自動車ソフトウェアプロセス改善・能力判定モデル'],
            ['SoT', 'Single Source of Truth（唯一の正）— project.json が該当'],
            ['フェーズゲート', '次フェーズへの移行可否を判定するチェックポイント'],
            ['セッション', 'Claude Code の 1 回の対話実行単位'],
        ],
        col_widths=[3.5, 12.5]
    )

    doc.add_heading('1.4 設計原則', level=2)
    add_styled_table(doc,
        ['原則ID', '原則', '説明'],
        [
            ['P-01', '自動検証可能性', '全ルールに自動検証手段を持たせる'],
            ['P-02', 'Single Source of Truth', 'project.json を数値・状態の唯一の正とする'],
            ['P-03', '静的解析連携', '型名・関数名を静的解析で自動検証する'],
            ['P-04', '自動クロスチェック', '文書とコード間の数値・名称を自動照合する'],
            ['P-05', 'ゲート強制', '全フェーズにゲート基準を設け、PASS まで次フェーズをブロックする'],
            ['P-06', 'バックアップ・ロールバック', '重要変更前にバックアップし、常にロールバック可能とする'],
            ['P-07', '自己プロセス適用', 'フレームワーク開発自体もこのプロセスに従い記録を残す'],
        ],
        col_widths=[2, 4, 10]
    )

    doc.add_page_break()

    # ========== 2. アーキテクチャ ==========
    doc.add_heading('2. フレームワークアーキテクチャ', level=1)

    doc.add_heading('2.1 リポジトリ構成', level=2)
    doc.add_paragraph(
        'TORASAN は全スキルの開発・メンテナンスの拠点であり、プロジェクトテンプレートを提供する。'
        '個別プロジェクト（機能安全 PJ、Web PJ 等）は独立リポジトリとして運用する。'
    )

    _insert_diagram(doc, draw_repo_structure, caption='図 2-1: リポジトリ構成',
                    dml_key='repo')
    doc.add_paragraph()

    add_styled_table(doc,
        ['パス', '役割', '管理対象'],
        [
            ['.claude/skills/', 'スキル格納',
             f'汎用 {SKILL_COUNT_UNIVERSAL}本 + ドメイン {SKILL_COUNT_DOMAIN}本 = 計 {SKILL_COUNT_TOTAL}本'],
            ['.claude/knowledge/', 'ナレッジ格納',
             f'汎用 {KNOWLEDGE_COUNT_UNIVERSAL}本 + ドメイン {KNOWLEDGE_COUNT_DOMAIN}本 = 計 {KNOWLEDGE_COUNT_TOTAL}本'],
            ['install.sh', 'スキル配布', '汎用スキル・ナレッジを ~/.claude/ に配布'],
            ['PROCESS.md', 'プロセス定義', f'ISO 26262 + SPICE V モデル（{PHASE_COUNT} フェーズ）'],
            ['project.json', 'プロジェクト SoT', '構成・進捗・変更ログ・TBD 管理'],
            ['docs/', '成果物格納', 'フェーズ成果物（安全計画、HARA 等）'],
            ['process_records/', 'SPICE エビデンス', 'プロセス実施記録・ゲートログ'],
        ],
        col_widths=[3.5, 4, 8.5]
    )

    doc.add_heading('2.2 二層配布モデル', level=2)
    doc.add_paragraph(
        'スキル配布は二層構造で運用する。汎用スキルは install.sh でグローバル配布し、'
        'ドメインスキルはプロジェクトレジストリ経由で選択配布する。'
    )

    _insert_diagram(doc, draw_distribution_model, caption='図 2-2: 二層配布モデル',
                    dml_key='dist')
    doc.add_paragraph()

    add_styled_table(doc,
        ['層', '対象', '配布手段', '配布先', 'スキル数'],
        [
            ['Tier 1（汎用）', 'セッション管理、監査、環境系', 'install.sh', '~/.claude/skills/',
             str(SKILL_COUNT_UNIVERSAL)],
            ['Tier 2（ドメイン）', '安全分析、コード生成、テスト系', '/repo-manage sync', 'PJ/.claude/skills/',
             str(SKILL_COUNT_DOMAIN)],
        ],
        col_widths=[3, 5, 3.5, 3.5, 2]
    )

    doc.add_heading('2.3 メモリ・状態管理アーキテクチャ', level=2)

    _insert_diagram(doc, draw_memory_architecture, caption='図 2-3: メモリ・状態管理アーキテクチャ',
                    dml_key='mem')
    doc.add_paragraph()

    add_styled_table(doc,
        ['ストレージ', 'パス', 'スコープ', 'ライフタイム'],
        [
            ['Auto Memory', '~/.claude/projects/{slug}/memory/MEMORY.md', 'セッション', '自動ロード（200 行上限）'],
            ['Session State', '~/.claude/projects/{slug}/memory/session_state.md', '引継ぎ', '/session end で上書き'],
            ['Project Knowledge', '.claude/knowledge/*.md', '永続（Git 管理）', 'ドメイン知識'],
            ['Process Records', 'process_records/*.md', '永続（Git 管理）', 'SPICE エビデンス'],
            ['Project Registry', '~/.claude/project_registry.json', 'グローバル', 'レジストリ情報'],
            ['Skill Manifest', '~/.claude/.shared-skills-manifest.json', 'グローバル', '配布履歴'],
        ],
        col_widths=[3, 6, 3, 4]
    )

    doc.add_page_break()

    # ========== 3. スキルカタログ ==========
    doc.add_heading('3. スキルカタログ', level=1)

    doc.add_heading('3.1 スキル分類体系', level=2)
    add_styled_table(doc,
        ['分類', '説明', 'スキル数'],
        [
            ['P1: コアプロセス', 'フェーズ実行・SPICE 評価・バリデーション', '4'],
            ['P2: 高頻度オペレーション', 'トレーサビリティ・ドキュメント・変更管理', '4'],
            ['D: ドメインエンジニアリング', '安全分析・組み込み設計・コード品質', '10'],
            ['E: 設計・テスト仕様', 'システム設計・要件生成・テスト設計', '6'],
            ['M: マネジメント・監査', '問題管理・構成監査・ヘルスチェック', '4'],
            ['O: オペレーション・環境', 'セッション・スキル管理・技術巡回', '9+'],
        ],
        col_widths=[4, 7, 2]
    )

    doc.add_heading('3.2 汎用スキル一覧（Tier 1 — グローバル配布）', level=2)
    add_styled_table(doc,
        ['#', 'スキル名', '機能概要'],
        [
            ['1', '/session', 'セッションライフサイクル管理（開始・終了・状態確認）'],
            ['2', '/dashboard', 'プロジェクトダッシュボード表示'],
            ['3', '/health-check', 'プロジェクト健全性評価'],
            ['4', '/skill-manage', 'スキル一覧・監査・スキャフォールド'],
            ['5', '/skill-evolve', 'スキル成熟度 PDCA サイクル管理'],
            ['6', '/repo-manage', 'プロジェクトレジストリ・ドメインスキル同期'],
            ['7', '/config-audit', '構成管理監査（SUP.8）'],
            ['8', '/update-record', 'SPICE プロセス記録更新'],
            ['9', '/env-check', '開発環境検証'],
            ['10', '/platform-info', 'プラットフォーム情報検出'],
            ['11', '/claude-master', '公式ドキュメント巡回・技術更新'],
            ['12', '/commit-change', 'Git コミット + 変更ログ記録'],
            ['13', '/backup', 'Git タグによるバックアップ・復元'],
            ['14', '/worktree-cleanup', 'ワークツリー整理'],
        ],
        col_widths=[1, 3.5, 11.5]
    )

    doc.add_heading('3.3 ドメインスキル一覧（Tier 2 — レジストリ配布）', level=2)
    add_styled_table(doc,
        ['#', 'スキル名', '機能概要', '分類'],
        [
            ['1', '/execute-phase', 'V モデルフェーズ実行（PH-01〜PH-15）', 'P1'],
            ['2', '/assess-spice', 'SPICE 能力レベル評価', 'P1'],
            ['3', '/validate', '成果物整合性検証（T-01〜T-10）', 'P1'],
            ['4', '/trace', 'トレーサビリティマトリクス生成', 'P2'],
            ['5', '/manage-tbd', 'TBD 項目管理', 'P2'],
            ['6', '/fmea', 'FMEA（故障モード影響分析）', 'D'],
            ['7', '/safety-diag', '安全診断カバレッジ設計', 'D'],
            ['8', '/safety-concept', '機能安全コンセプト設計', 'D'],
            ['9', '/safety-verify', '安全ケース検証（GSN）', 'D'],
            ['10', '/mcu-config', 'MCU ペリフェラル設定コード生成', 'D'],
            ['11', '/memory-map', 'メモリマップ設計', 'D'],
            ['12', '/driver-gen', 'デバイスドライバコード生成', 'D'],
            ['13', '/motor-control', 'BLDC モーター制御コード生成', 'D'],
            ['14', '/static-analysis', '静的解析（MISRA C:2012）', 'D'],
            ['15', '/hw-review', 'ハードウェア設計レビュー', 'D'],
            ['16', '/system-design', 'システムアーキテクチャ設計', 'E'],
            ['17', '/sw-design', 'ソフトウェア詳細設計', 'E'],
            ['18', '/srs-generate', 'ソフトウェア要件仕様書生成', 'E'],
            ['19', '/test-design', 'ユニットテスト仕様設計', 'E'],
            ['20', '/test-coverage', 'テストカバレッジ分析', 'E'],
            ['21', '/systest-design', 'システム・統合テスト仕様', 'E'],
            ['22', '/problem-resolve', '問題解決管理（SUP.9）', 'M'],
            ['23', '/select-standard', '適用規格選択ガイド', 'O'],
            ['24', '/switch-standard', '規格構成切替', 'O'],
            ['25', '/generate-docs', '文書生成（docx/xlsx）', 'O'],
            ['26', '/ingest', 'レガシー資産取込', 'O'],
        ],
        col_widths=[1, 3.5, 8.5, 1.5]
    )

    doc.add_heading('3.4 スキル成熟度モデル', level=2)
    add_styled_table(doc,
        ['レベル', 'ステージ', '判定基準', '遷移条件'],
        [
            ['L1', 'Draft', 'SKILL.md が存在する', '自動'],
            ['L2', 'Active', 'ナレッジ参照セクションがある', '手動レビュー'],
            ['L3', 'Practiced', 'フィードバック記録が 1 件以上', 'データ存在で自動'],
            ['L4', 'Mature', '改善 1 回以上 + 平均評価 >= 4.0', 'アセスメント'],
            ['L5', 'Optimized', '改善 2 回以上 + 評価 >= 4.5 + ナレッジ更新済み', 'アセスメント'],
        ],
        col_widths=[2, 2.5, 6.5, 5]
    )

    doc.add_page_break()

    # ========== 4. ナレッジベース ==========
    doc.add_heading('4. ナレッジベース仕様', level=1)

    doc.add_heading('4.1 グローバルナレッジ（全 PJ 共通）', level=2)
    doc.add_paragraph('配布先: ~/.claude/knowledge/')
    add_styled_table(doc,
        ['#', 'ファイル名', '内容'],
        [
            ['1', 'claude_code_ops.md', 'Claude Code 操作手順（セッション管理、メモリシステム）'],
            ['2', 'claude_platform_updates.md', 'プラットフォーム機能追跡・変更ログ監視'],
            ['3', 'skill_lifecycle.md', 'スキル成熟度モデル（L1〜L5 定義）'],
            ['4', 'skill_feedback_log.md', 'スキルフィードバック収集・改善追跡'],
            ['5', 'git_worktree_branch_management.md', 'Git ワークツリー操作手順'],
            ['6', 'cross_platform_dev.md', 'マルチプラットフォーム開発（パス変換等）'],
            ['7', 'memory_paths.md', 'メモリ・状態ファイルパス構成'],
            ['8', 'error_prevention.md', 'エラー防止・PDCA プロセス定義'],
        ],
        col_widths=[1, 6, 9]
    )

    doc.add_heading('4.2 ドメインナレッジ（PJ 固有）', level=2)
    doc.add_paragraph('格納先: .claude/knowledge/')
    add_styled_table(doc,
        ['#', 'ファイル名', '内容'],
        [
            ['1', 'automotive_spice.md', 'SPICE 評価リファレンス（PA 評価、能力レベル）'],
            ['2', 'bldc_safety.md', 'BLDC モーター安全分析'],
            ['3', 'fmea_guide.md', 'FMEA 手法ガイド（重大度・発生度・検出度）'],
            ['4', 'iso26262_iec60730.md', 'デュアルコンプライアンスリファレンス'],
            ['5', 'misra_c_2012.md', 'MISRA-C:2012 コーディング規約'],
            ['6', 'product_standard_mapping.md', '規格選定ガイド（自動車/家電/産業/医療）'],
            ['7', 'safety_case_gsn.md', 'GSN による安全ケース記述'],
            ['8', 'safety_diagnostics.md', '診断カバレッジ・SFF 計算手法'],
            ['9', 'srs_template.md', 'ソフトウェア要件仕様書テンプレート'],
            ['10', 'pptx_advanced_shapes.md', 'python-pptx 高度図形操作リファレンス'],
            ['11', 'uml_diagramming.md', 'UML 作図リファレンス（python-pptx 実装パターン）'],
        ],
        col_widths=[1, 6, 9]
    )

    doc.add_page_break()

    # ========== 5. プロセス定義 ==========
    doc.add_heading('5. プロセス定義', level=1)

    doc.add_heading('5.1 V モデルフェーズ定義', level=2)
    doc.add_paragraph(
        f'TORASAN は ISO 26262 に基づく {PHASE_COUNT} フェーズの V モデルを採用する。'
        '各フェーズは INPUT/OUTPUT/PROCEDURE/GATE で構成される。'
    )

    _insert_diagram(doc, draw_vmodel, width=Inches(6.5),
                    caption=f'図 5-1: V モデル {PHASE_COUNT} フェーズ構成',
                    dml_key='vmodel')
    doc.add_paragraph()

    add_styled_table(doc,
        ['ID', 'フェーズ名', 'SPICE プロセス', '成果物'],
        [
            ['PH-01', 'プロジェクト計画・安全計画', 'MAN.3', '安全計画書'],
            ['PH-02', 'アイテム定義', 'SYS.1', 'アイテム定義書'],
            ['PH-03', 'HARA', 'SYS.1', 'HARA シート・安全目標'],
            ['PH-04', 'FSC（機能安全コンセプト）', 'SYS.2', 'FSC 文書'],
            ['PH-05', 'TSC（技術安全コンセプト）', 'SYS.3', 'TSC 文書'],
            ['PH-06', 'システム設計', 'SYS.3', 'システムアーキテクチャ文書'],
            ['PH-07', 'HW 設計', 'SYS.3', 'HW 設計書・FMEDA'],
            ['PH-08', 'SW 安全要求（SRS）', 'SWE.1', 'SRS 文書'],
            ['PH-09', 'SW アーキテクチャ設計', 'SWE.2', 'SW アーキテクチャ文書'],
            ['PH-10', 'SW ユニット設計・実装', 'SWE.3', 'ソースコード・設計書'],
            ['PH-11', 'SW テスト仕様・実施', 'SWE.4-6', 'テスト仕様書・結果'],
            ['PH-12', 'HW 統合テスト', 'SYS.4', '統合テスト結果'],
            ['PH-13', 'システムテスト', 'SYS.5', 'システムテスト結果'],
            ['PH-14', '機能安全検証', 'SYS.5', '安全検証レポート'],
            ['PH-15', '安全アセスメント', 'SUP.1', '安全アセスメントレポート'],
        ],
        col_widths=[2, 5.5, 3.5, 5]
    )

    doc.add_heading('5.2 Automotive SPICE プロセス定義', level=2)
    add_styled_table(doc,
        ['プロセス ID', 'プロセス名', 'カテゴリ'],
        [
            ['SYS.1', '要件導出', 'システムエンジニアリング'],
            ['SYS.2', 'システム要件分析', 'システムエンジニアリング'],
            ['SYS.3', 'システムアーキテクチャ設計', 'システムエンジニアリング'],
            ['SYS.4', 'システム統合・統合テスト', 'システムエンジニアリング'],
            ['SYS.5', 'システム適格性テスト', 'システムエンジニアリング'],
            ['SWE.1', 'SW 要件分析', 'ソフトウェアエンジニアリング'],
            ['SWE.2', 'SW アーキテクチャ設計', 'ソフトウェアエンジニアリング'],
            ['SWE.3', 'SW 詳細設計・ユニット構築', 'ソフトウェアエンジニアリング'],
            ['SWE.4', 'SW ユニット検証', 'ソフトウェアエンジニアリング'],
            ['SWE.5', 'SW 統合・統合テスト', 'ソフトウェアエンジニアリング'],
            ['SWE.6', 'SW 適格性テスト', 'ソフトウェアエンジニアリング'],
            ['MAN.3', 'プロジェクト管理', 'マネジメント'],
            ['SUP.1', '品質保証', 'サポート'],
            ['SUP.8', '構成管理', 'サポート'],
            ['SUP.9', '問題解決管理', 'サポート'],
            ['SUP.10', '変更要求管理', 'サポート'],
        ],
        col_widths=[3, 6, 7]
    )

    doc.add_heading('5.3 SPICE 能力レベル判定基準', level=2)
    add_styled_table(doc,
        ['レベル', '名称', 'PA 1.1', 'PA 2.1', 'PA 2.2'],
        [
            ['Level 0', '不完全', 'N or P', '—', '—'],
            ['Level 1', '実施済み', 'L or F', '—', '—'],
            ['Level 2', '管理済み', 'F', 'L or F', 'L or F'],
        ],
        col_widths=[2, 3, 3.5, 3.5, 3.5]
    )

    doc.add_paragraph()
    doc.add_paragraph('評価スケール:')
    for label, desc in [
        ('N (Not achieved)', '0〜15%'),
        ('P (Partially achieved)', '16〜50%'),
        ('L (Largely achieved)', '51〜85%'),
        ('F (Fully achieved)', '86〜100%'),
    ]:
        doc.add_paragraph(f'{label}: {desc}', style='List Bullet')

    doc.add_page_break()

    # ========== 6. データ構造 ==========
    doc.add_heading('6. データ構造仕様', level=1)

    doc.add_heading('6.1 project.json スキーマ', level=2)
    doc.add_paragraph(
        'project.json はプロジェクトの唯一の正（Single Source of Truth）として機能する。'
        'スキーマバージョン 3.0 に準拠する。'
    )
    add_styled_table(doc,
        ['フィールド', '型', '説明'],
        [
            ['$schema', 'string', 'スキーマバージョン（"3.0"）'],
            ['projectName', 'string', 'プロジェクト名'],
            ['category', 'string', 'カテゴリ（開発基盤、組み込み開発 等）'],
            ['mode', 'enum', '動作モード（EXPLORE / DEVELOP / REVIEW）'],
            ['description', 'string', 'プロジェクト概要'],
            ['owner', 'string', 'オーナー名'],
            ['standard.base', 'string', '基本規格（ISO 26262:2018 等）'],
            ['standard.product_override', 'string', '製品規格（IEC 60730 等）'],
            ['standard.product_override_class', 'string', '製品分類（Class B 等）'],
            ['phases[]', 'array', 'フェーズ定義（id, name, status, doc_version, gates）'],
            ['tbd_items[]', 'array', 'TBD 項目（id, description, dependency）'],
            ['changeLog[]', 'array', '変更ログ（id, date, type, item, from, to, reason）'],
            ['spice_assessment', 'object', 'SPICE 評価結果（target_level, processes）'],
            ['traceability', 'object', 'トレーサビリティ情報（SG, FSR, TSR, SR, TC）'],
            ['problemLog[]', 'array', '問題ログ（id, severity, title, status）'],
            ['reviewLog[]', 'array', 'レビューログ（id, phase, reviewer, findings）'],
            ['gateLog[]', 'array', 'ゲートログ（phase, gates, result, reviewer）'],
        ],
        col_widths=[4, 2, 10]
    )

    doc.add_heading('6.2 session_state.md 構造', level=2)
    doc.add_paragraph('セッション引継ぎ情報を格納する揮発性ファイル。/session end で上書き保存される。')
    add_styled_table(doc,
        ['セクション', '内容', '更新タイミング'],
        [
            ['直前の作業', '作業内容・到達点・次のアクション', '/session end'],
            ['Git 状態', 'ブランチ・最新コミット・未コミット変更・master 差分', '/session end'],
            ['プロジェクト状態', 'フェーズ進捗・TBD 数・未解決問題', '/session end'],
            ['重要なコンテキスト', '反省会で得た教訓・注意事項', '/session end'],
            ['未完了タスク', '反省会の「未」改善項目', '/session end'],
        ],
        col_widths=[4, 7, 5]
    )

    doc.add_heading('6.3 プロジェクトレジストリ構造', level=2)
    doc.add_paragraph('格納先: ~/.claude/project_registry.json')
    add_styled_table(doc,
        ['フィールド', '型', '説明'],
        [
            ['projects[].name', 'string', 'プロジェクト識別名'],
            ['projects[].path', 'string', '絶対パス'],
            ['projects[].category', 'enum', 'development | research | maintenance'],
            ['projects[].last_synced', 'ISO8601', '最終同期日時'],
            ['projects[].sync_status', 'enum', 'success | pending | error'],
            ['projects[].skill_sync_count', 'integer', '同期済みドメインスキル数'],
        ],
        col_widths=[4.5, 2.5, 9]
    )

    doc.add_page_break()

    # ========== 7. 安全規格サポート ==========
    doc.add_heading('7. 安全規格サポート', level=1)

    doc.add_heading('7.1 対応規格一覧', level=2)
    add_styled_table(doc,
        ['規格', '版数', '対象ドメイン', 'TORASAN 対応状態'],
        [
            ['ISO 26262', '2018', '自動車 E/E システム', '対応済み（プロセスフレームワーク）'],
            ['IEC 60730-1', 'Class A〜C', '家電制御機器', '対応済み（製品分類）'],
            ['IEC 61508', 'SIL 1〜4', '産業用 E/E/PE', 'テンプレート構造のみ'],
            ['IEC 62304', '—', '医療機器ソフトウェア', '未対応'],
            ['EN 50128', '—', '鉄道ソフトウェア', '未対応'],
        ],
        col_widths=[3, 3, 5, 5]
    )

    doc.add_heading('7.2 規格選択ルール', level=2)
    doc.add_paragraph(
        'project.json の standard セクションで規格を選択する。'
        '規格情報は standard セクションに集約し、実装コードは各PJリポの src/ に配置する。'
    )
    add_styled_table(doc,
        ['カテゴリ', 'product_override', 'product_override_class'],
        [
            ['車載', 'null', '—'],
            ['家電', 'IEC60730', 'ClassB'],
            ['産業', 'IEC61508', 'SIL2'],
        ],
        col_widths=[5, 7, 4]
    )

    doc.add_heading('7.3 安全メトリクス', level=2)
    add_styled_table(doc,
        ['メトリクス', '目標値', '算出式'],
        [
            ['SPFM（単一故障メトリクス）', '>= 90%', '1 - (lambda_SPF / lambda_total)'],
            ['LFM（潜在故障メトリクス）', '>= 60%', '1 - (lambda_MPF_latent / lambda_total)'],
            ['PMHF（HW 故障確率）', '< 1E-7/h', '残留故障率の合計'],
            ['DC（診断カバレッジ）', '>= 90%（Medium）', '検出故障 / 全故障'],
        ],
        col_widths=[5, 3.5, 7.5]
    )

    doc.add_page_break()

    # ========== 8. テストカバレッジ目標 ==========
    doc.add_heading('8. 品質目標', level=1)

    doc.add_heading('8.1 テストカバレッジ目標', level=2)
    add_styled_table(doc,
        ['カバレッジ種別', '目標値', '適用フェーズ'],
        [
            ['ステートメントカバレッジ', '>= 95%', 'PH-11'],
            ['ブランチカバレッジ', '>= 90%', 'PH-11'],
            ['MC/DC', '>= 85%', 'PH-11'],
            ['要件-テスト紐付け率', '100%', 'PH-11〜PH-14'],
        ],
        col_widths=[5, 3, 8]
    )

    doc.add_heading('8.2 トレーサビリティ目標', level=2)
    add_styled_table(doc,
        ['メトリクス', '目標値', '測定方法'],
        [
            ['要件リンク完全性', '100%', '/trace で孤立要件検出'],
            ['前方トレーサビリティ', '100%', 'SG→FSR→TSR→SR→TC 全チェーン'],
            ['後方トレーサビリティ', '100%', 'TC→SR→TSR→FSR→SG 逆引き'],
            ['カバレッジギャップ', '0%', '/validate T-01〜T-10 ルール'],
        ],
        col_widths=[5, 3, 8]
    )

    doc.add_heading('8.3 コーディング規約', level=2)
    doc.add_paragraph('全 C コードに MISRA-C:2012 を適用する。静的解析ツールチェーン:')
    for tool, desc in [
        ('cppcheck', '汎用静的解析（メモリリーク、未初期化変数等）'),
        ('clang-tidy', 'LLVM ベースリンター（コードスタイル、パターン検出）'),
        ('flawfinder', 'セキュリティ脆弱性スキャナ'),
    ]:
        doc.add_paragraph(f'{tool}: {desc}', style='List Bullet')

    doc.add_heading('8.4 開発環境ツール管理', level=2)
    doc.add_paragraph(
        'TORASAN の GUI 管理アプリ（localhost:3000）は環境監査機能を備え、'
        '全ツールのバージョン・パス・導入状態を自動検査する。'
        'Windows 環境では cmd.exe の PATH と Git Bash の PATH が異なるため、'
        'KNOWN_TOOL_PATHS による既知パス解決メカニズムを実装している。'
    )

    _insert_diagram(doc, draw_tool_resolution, width=Inches(5.5),
                    caption='図 8-1: ツールパス解決フロー',
                    dml_key='tool')
    doc.add_paragraph()

    add_styled_table(doc,
        ['ツール', 'カテゴリ', 'パス解決方法', 'バージョン'],
        [
            ['git', '必須', 'PATH（cmd.exe）', '2.53+'],
            ['python', '必須', 'PATH（cmd.exe）', '3.14+'],
            ['node', '必須', 'PATH（cmd.exe）', '24.x LTS'],
            ['npm', '必須', 'PATH（cmd.exe）', '11.x'],
            ['cppcheck', '静的解析', 'KNOWN_TOOL_PATHS', '2.20+'],
            ['clang-tidy', '静的解析', 'KNOWN_TOOL_PATHS', 'LLVM 21+'],
            ['flawfinder', '静的解析', 'KNOWN_TOOL_PATHS（venv）', '2.0+'],
            ['gh', 'GitHub', 'KNOWN_TOOL_PATHS', '2.87+'],
        ],
        col_widths=[3, 2.5, 5, 2.5]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'ツール追加時は error_prevention.md §E「ツール追加時の基本フロー」に従い、'
        'KNOWN_TOOL_PATHS 登録 → ビルド確認 → GUI 検証 → env_state.md 更新 → '
        '仕様書再生成（python scripts/generate_manuals.py）を一貫して実行する。'
    )

    doc.add_page_break()

    # ========== 9. 継続的改善 ==========
    doc.add_heading('9. 継続的改善プロセス', level=1)

    doc.add_heading('9.1 技術巡回サイクル', level=2)
    add_styled_table(doc,
        ['頻度', 'アクション', '担当スキル'],
        [
            ['セッション開始時', '7 日以上経過時に scan 提案', '/claude-master'],
            ['スキル実行後', 'フィードバック記録', '/skill-evolve'],
            ['5 セッション毎', '品質監査', '/skill-manage audit'],
        ],
        col_widths=[4, 7, 5]
    )

    doc.add_heading('9.2 反省会プロセス', level=2)
    doc.add_paragraph(
        '各セッション終了時に反省会（レトロスペクティブ）を実施し、'
        '議事録を process_records/retrospective/ に保存する。'
    )
    doc.add_paragraph('議事録構成:')
    sections = [
        'セッション成果サマリ',
        '発生した問題',
        '根本原因分析',
        '良かった点',
        '改善すべき点（対策・状態「済/未」付き）',
        '教訓',
        '前回反省会フォローアップ',
    ]
    for i, s in enumerate(sections, 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')

    # ========== 10. マニュアル管理システム ==========
    doc.add_page_break()
    doc.add_heading('10. マニュアル管理システム', level=1)

    doc.add_heading('10.1 概要', level=2)
    doc.add_paragraph(
        'TORASAN フレームワークの仕様書・操作手順書は Python スクリプトによるコード生成方式で管理する。'
        'フレームワークの変更に追従して文書を再生成し、差分を自動検出・記録する仕組みを備える。'
    )

    doc.add_heading('10.2 構成ファイル', level=2)
    add_styled_table(doc,
        ['ファイル', 'パス', '役割'],
        [
            ['統合生成スクリプト', 'scripts/generate_manuals.py', '再生成 + 差分管理のエントリポイント'],
            ['仕様書生成', 'scripts/generate_spec_doc.py', '仕様書の内容定義'],
            ['操作手順書生成', 'scripts/generate_ops_manual.py', '操作手順書の内容定義'],
            ['共通ユーティリティ', 'scripts/docx_utils.py', 'テーブル・スタイル等の共通関数'],
            ['マニフェスト', 'docs/manuals/.manuals_manifest.json', 'バージョン・ハッシュ・履歴'],
            ['差分レポート', 'docs/manuals/diff_reports/', '再生成ごとの差分記録'],
        ],
        col_widths=[4, 6, 6]
    )

    doc.add_heading('10.3 マニフェスト構造', level=2)
    doc.add_paragraph(
        '.manuals_manifest.json は文書のバージョン管理・変更履歴を一元管理する。'
    )
    add_styled_table(doc,
        ['フィールド', '型', '説明'],
        [
            ['version', 'integer', '全体バージョン番号（最大値）'],
            ['last_updated', 'string', '最終更新日時'],
            ['documents.{key}.version', 'integer', '文書別バージョン番号'],
            ['documents.{key}.hash', 'string', 'SHA-256 ハッシュ'],
            ['documents.{key}.generated_at', 'string', '生成日時'],
            ['documents.{key}.file', 'string', '出力ファイル名'],
            ['history[].version', 'integer', '履歴バージョン'],
            ['history[].changes.added', 'integer', '追加行数'],
            ['history[].changes.removed', 'integer', '削除行数'],
            ['history[].changes.sections', 'integer', '変更セクション数'],
            ['history[].diff_report', 'string', '差分レポートファイルパス'],
        ],
        col_widths=[5.5, 2, 8.5]
    )

    doc.add_heading('10.4 差分検出メカニズム', level=2)
    doc.add_paragraph('再生成時の差分検出は以下の手順で実行される:')
    steps = [
        '既存 docx からプレーンテキストを抽出（段落 + テーブルセル）',
        'コンテンツ生成スクリプトを実行して新 docx を生成',
        '新 docx からプレーンテキストを抽出',
        'unified diff アルゴリズムで差分を計算',
        '差分レポートを docs/manuals/diff_reports/ に保存',
        'マニフェストにバージョン番号・ハッシュ・差分統計を記録',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')

    doc.add_heading('10.5 更新トリガーと運用フロー', level=2)
    add_styled_table(doc,
        ['トリガー', '実行タイミング', '操作'],
        [
            ['スキル追加・変更', 'スキル定義変更後', 'python scripts/generate_manuals.py'],
            ['プロセス定義変更', 'PROCESS.md 更新後', 'python scripts/generate_manuals.py'],
            ['フレームワーク構成変更', 'CLAUDE.md / project.json 変更後', 'python scripts/generate_manuals.py'],
            ['セッション終了時', '/session end フロー内', '自動提案（任意実行）'],
        ],
        col_widths=[4, 5, 7]
    )

    # ========== 付録 ==========
    doc.add_page_break()
    doc.add_heading('付録A: 変更管理エントリ構造', level=1)
    add_styled_table(doc,
        ['フィールド', '型', '説明'],
        [
            ['id', 'string', '変更ID（CHG-NNN）'],
            ['date', 'string', '変更日（YYYY-MM-DD）'],
            ['type', 'enum', 'プロセス改善 | 要件変更 | バグ修正 | その他'],
            ['item', 'string', '影響を受けるファイル'],
            ['from', 'string', '変更前の状態'],
            ['to', 'string', '変更後の状態'],
            ['reason', 'string', '変更理由'],
            ['affectedPhases', 'array', '影響フェーズ（PH-XX）'],
            ['backupRef', 'string', 'バックアップ参照（archive/vX/ or git tag）'],
            ['approvedBy', 'string', '承認者'],
            ['status', 'enum', '完了 | 進行中 | 保留中'],
        ],
        col_widths=[4, 2, 10]
    )

    # 保存
    doc.save(output_path)
    print(f'仕様書を保存しました: {output_path}')


if __name__ == '__main__':
    output = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        'docs', 'manuals', 'TORASAN_フレームワーク仕様書.docx'
    )
    create_spec_doc(output)
